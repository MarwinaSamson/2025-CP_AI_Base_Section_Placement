from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_http_methods
from django.template.loader import render_to_string
from xhtml2pdf import pisa
import json
import io
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from enrollment_app.models import ProgramSelection
from admin_app.models import Section, SchoolYear
from coordinator_app.models import Qualified_for_ste

from django.views.decorators.csrf import csrf_exempt
from coordinator_app.models import AIAssistantPreference
from admin_app.models import Program


@login_required
def section_assignment(request):
    """Section assignment dashboard scoped to the coordinator's program."""

    user_profile = getattr(request.user, 'profile', None)
    program_code = user_profile.program.code if user_profile and user_profile.program else None
    program_name = user_profile.program.name if user_profile and user_profile.program else None

    user_full_name = request.user.get_full_name() or request.user.username
    user_type = f"{program_code} Coordinator" if program_code else "Coordinator"
    user_photo = user_profile.photo.url if user_profile and user_profile.photo else None

    name_parts = user_full_name.split()
    user_initials = ''.join(part[0].upper() for part in name_parts[:2]) if name_parts else "CO"

    students_payload = []
    sections_payload = []

    if program_code:
        # Active school year (fallback to latest if none marked active)
        active_sy = (
            SchoolYear.objects.filter(is_active=True).order_by('-start_date').first()
            or SchoolYear.objects.order_by('-start_date').first()
        )

        # Sections for this program and active school year
        sections_qs = Section.objects.filter(program__code=program_code)
        if active_sy:
            sections_qs = sections_qs.filter(school_year=active_sy)

        # Update counts from database before building payload
        for section in sections_qs:
            section.update_current_students_count()

        sections_payload = [
            {
                'id': str(section.id),
                'name': section.name,
                'capacity': section.max_students,
                'current': section.current_students,
            }
            for section in sections_qs
        ]

        selections = (
            ProgramSelection.objects
            .select_related('student', 'student__student_data')
            .filter(selected_program_code=program_code)
        )

        lrns = [sel.student.lrn for sel in selections]
        score_map = {
            rec.student_lrn: rec
            for rec in Qualified_for_ste.objects.filter(student_lrn__in=lrns)
        }

        for sel in selections:
            student = sel.student
            student_data = getattr(student, 'student_data', None)

            name_parts = [
                getattr(student_data, 'last_name', ''),
                getattr(student_data, 'first_name', ''),
                getattr(student_data, 'middle_name', '') or ''
            ]
            display_name = ', '.join(
                [name_parts[0], ' '.join(name_parts[1:]).strip()]
            ).strip(', ')

            scores = score_map.get(student.lrn)
            exam_score = float(scores.exam_score) if scores and scores.exam_score else 0
            interview_score = float(scores.interview_score) if scores and scores.interview_score else 0

            students_payload.append({
                'name': display_name or student.lrn,
                'lrn': student.lrn,
                'exam': exam_score,
                'interview': interview_score,
                'finalSection': sel.assigned_section or None,
                'admin_approved': sel.admin_approved,
            })

    # Check if AI Assistant is enabled for this coordinator and program
    ai_enabled = False
    if program_code:
        try:
            program = user_profile.program
            ai_pref = AIAssistantPreference.objects.filter(
                user=request.user,
                program=program,
                ai_enabled=True
            ).first()
            ai_enabled = ai_pref is not None
        except Exception:
            ai_enabled = False

    context = {
        'program_code': program_code,
        'program_name': program_name,
        'students_json': json.dumps(students_payload),
        'sections_json': json.dumps(sections_payload),
        'user_full_name': user_full_name,
        'user_type': user_type,
        'user_photo': user_photo,
        'user_initials': user_initials,
        'ai_enabled': ai_enabled,  # ADD THIS
    }

    return render(request, 'coordinator_app/sectionAssignment.html', context)


@login_required
@require_http_methods(["POST"])
def export_assignments_pdf(request):
    """Export section assignments as PDF (Windows-safe)"""

    user_profile = getattr(request.user, 'profile', None)
    program_code = user_profile.program.code if user_profile and user_profile.program else "N/A"
    program_name = user_profile.program.name if user_profile and user_profile.program else "Unknown Program"

    try:
        data = json.loads(request.body)
        students = data.get('students', [])

        context = {
            'program_code': program_code,
            'program_name': program_name,
            'students': students,
            'generated_date': datetime.now().strftime('%B %d, %Y'),
            'generated_time': datetime.now().strftime('%I:%M %p'),
            'total_students': len(students),
        }

        html = render_to_string(
            'coordinator_app/exports/section_assignment_pdf.html',
            context
        )

        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = (
            f'attachment; filename="Section_Assignments_'
            f'{program_code}_{datetime.now().strftime("%Y%m%d")}.pdf"'
        )

        pisa_status = pisa.CreatePDF(
            html,
            dest=response,
            encoding='UTF-8'
        )

        if pisa_status.err:
            return JsonResponse({'error': 'PDF generation failed'}, status=500)

        return response

    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)
    
@login_required
@require_http_methods(["POST"])
def export_assignments_docx(request):
    """Export section assignments as DOCX"""

    user_profile = getattr(request.user, 'profile', None)
    program_code = user_profile.program.code if user_profile and user_profile.program else "N/A"
    program_name = user_profile.program.name if user_profile and user_profile.program else "Unknown Program"

    try:
        data = json.loads(request.body)
        students = data.get('students', [])

        doc = Document()

        # Page margins
        section = doc.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        # Title
        title = doc.add_heading('Section Assignment Report', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(153, 27, 27)

        # Program info
        info = doc.add_paragraph(f"{program_name} ({program_code})")
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_info = doc.add_paragraph(
            f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
        )
        date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'

        headers = [
            'Student Name',
            'LRN',
            'Exam Score',
            'Interview Score',
            'Final Section'
        ]

        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '991b1b')
            cell._element.get_or_add_tcPr().append(shading)

        for student in students:
            row = table.add_row().cells
            row[0].text = student.get('name', '')
            row[1].text = student.get('lrn', '')
            row[2].text = f"{student.get('exam', 0)}%"
            row[3].text = f"{student.get('interview', 0)}%"
            row[4].text = student.get('finalSection', '-')

            for cell in row:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = (
            f'attachment; filename="Section_Assignments_'
            f'{program_code}_{datetime.now().strftime("%Y%m%d")}.docx"'
        )

        return response

    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@login_required
@require_http_methods(["POST"])
def toggle_ai_mode(request):
    """Toggle AI automation on/off for the coordinator's program"""
    
    try:
        data = json.loads(request.body)
        ai_enabled = data.get('ai_enabled', False)
        program_code = data.get('program_code')
        
        if not program_code:
            return JsonResponse({'error': 'Program code is required'}, status=400)
        
        # Get the program
        try:
            program = Program.objects.get(code=program_code)
        except Program.DoesNotExist:
            return JsonResponse({'error': 'Program not found'}, status=404)
        
        # Update or create AI preference
        ai_pref, created = AIAssistantPreference.objects.update_or_create(
            user=request.user,
            program=program,
            defaults={'ai_enabled': ai_enabled}
        )
        
        action = 'created' if created else 'updated'
        status_text = 'enabled' if ai_enabled else 'disabled'
        
        return JsonResponse({
            'success': True,
            'message': f'AI automation {status_text} successfully',
            'ai_enabled': ai_pref.ai_enabled,
            'action': action
        })
        
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON data'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)