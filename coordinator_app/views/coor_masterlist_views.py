from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse
from django.contrib import messages
from django.db import transaction
from django.db.models import Avg, Count, Q
from decimal import Decimal, InvalidOperation
from admin_app.decorators import coordinator_required
from admin_app.models import Section, Subject, Program
from enrollment_app.models import Student, ProgramSelection
from coordinator_app.models import AcademicPerformance, GradeUploadBatch
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv
import io


@coordinator_required
def masterlist_by_section(request, section_id):
    """View to display masterlist for a specific section"""
    try:
        user_profile = request.user.profile
        program = user_profile.program

        section = get_object_or_404(
            Section.objects.select_related('program', 'school_year', 'adviser'),
            id=section_id,
            program=program
        )

        program_selections = ProgramSelection.objects.filter(
            assigned_section_id=section_id,
            admin_approved=True
        ).select_related(
            'student',
            'student__student_data',
        ).order_by('student__student_data__last_name', 'student__student_data__first_name')

        students_data = []
        for idx, ps in enumerate(program_selections, 1):
            student = ps.student
            student_info = getattr(student, 'student_data', None)

            if student_info:
                students_data.append({
                    'number': idx,
                    'lrn': student.lrn,
                    'last_name': student_info.last_name,
                    'first_name': student_info.first_name,
                    'middle_name': student_info.middle_name or '',
                    'full_name': student_info.full_name,
                    'gender': student_info.gender,
                    'age': student_info.age,
                    'enrolling_as': student_info.enrolling_as if student_info.enrolling_as else [],
                    'status': 'Enrolled' if ps.admin_approved else 'Pending',
                })

        total_students = len(students_data)
        male_count = sum(1 for s in students_data if s['gender'] == 'male')
        female_count = sum(1 for s in students_data if s['gender'] == 'female')

        ages = [s['age'] for s in students_data if s['age'] is not None]
        average_age = sum(ages) / len(ages) if ages else 0

        slots_remaining = section.max_students - total_students

        context = {
            'user': request.user,
            'section': section,
            'students': students_data,
            'program': program.code if program else '',
            'program_full_name': program.name if program else '',
            'program_code': program.code if program else '',
            'user_profile': user_profile,
            'total_students': total_students,
            'male_count': male_count,
            'female_count': female_count,
            'male_percentage': (male_count / total_students * 100) if total_students > 0 else 0,
            'female_percentage': (female_count / total_students * 100) if total_students > 0 else 0,
            'average_age': round(average_age, 1),
            'slots_remaining': slots_remaining,
            'active_school_year': section.school_year,
        }
        return render(request, 'coordinator_app/cor-masterlist.html', context)
    except Exception as e:
        context = {
            'user': request.user,
            'error': str(e),
            'students': [],
        }
        return render(request, 'coordinator_app/cor-masterlist.html', context)


# ─────────────────────────────────────────────────────────────
# IMPORT — Download Template (dynamic from Subject records)
# ─────────────────────────────────────────────────────────────

def _resolve_subject_program(section):
    """
    For REGULAR sections, subjects are stored under the track program (HETERO or TOP5).
    For all other programs (STE, SPFL, SPTVE, OHSP, etc.) use section.program directly.
    """
    if section.program and section.program.code == 'REGULAR' and section.regular_track:
        try:
            return Program.objects.get(code=section.regular_track)
        except Program.DoesNotExist:
            pass
    return section.program


@coordinator_required
def download_import_template(request, section_id):
    """Download a grade upload template whose columns are the program's active subjects."""
    user_profile = request.user.profile
    program = user_profile.program

    section = get_object_or_404(
        Section.objects.select_related('program', 'grade_level', 'school_year'),
        id=section_id,
        program=program,
    )

    subject_program = _resolve_subject_program(section)
    subjects = Subject.objects.filter(
        program=subject_program, is_active=True
    ).order_by('name')

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Grades'

    headers = ['LRN'] + [s.name for s in subjects]

    header_fill = PatternFill(start_color='991B1B', end_color='991B1B', fill_type='solid')
    header_font = Font(bold=True, color='FFFFFF')

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center')
        ws.column_dimensions[cell.column_letter].width = max(len(h) + 4, 14)

    ws.row_dimensions[1].height = 20

    # Sample row so coordinators know the expected format
    sample = ['123456789012'] + ['' for _ in subjects]
    for col, val in enumerate(sample, 1):
        cell = ws.cell(row=2, column=col, value=val)
        cell.font = Font(italic=True, color='9CA3AF')
        cell.alignment = Alignment(horizontal='center')

    grade_label = section.grade_level.code if section.grade_level else 'grades'
    prog_label = section.program.code if section.program else 'program'

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = (
        f'attachment; filename="grade_template_{prog_label}_{grade_label}.xlsx"'
    )
    wb.save(response)
    return response


# ─────────────────────────────────────────────────────────────
# IMPORT — Handle Upload (CSV / Excel) → AcademicPerformance
# ─────────────────────────────────────────────────────────────

@coordinator_required
def import_students(request, section_id):
    """Upload CSV/Excel grade sheet and save records to AcademicPerformance."""
    if request.method != 'POST':
        return redirect('coordinator:masterlist_by_section', section_id=section_id)

    user_profile = request.user.profile
    program = user_profile.program

    section = get_object_or_404(
        Section.objects.select_related('program', 'grade_level', 'school_year'),
        id=section_id,
        program=program,
    )

    # Validate quarter from the modal form
    quarter_str = request.POST.get('quarter', '').strip()
    if quarter_str not in ('1', '2', '3', '4', '5'):
        messages.error(request, 'Please select a valid period (Q1–Q4 or Final Grade) before uploading.')
        return redirect('coordinator:masterlist_by_section', section_id=section_id)
    quarter = int(quarter_str)

    import_file = request.FILES.get('import_file')
    if not import_file:
        messages.error(request, 'No file was uploaded.')
        return redirect('coordinator:masterlist_by_section', section_id=section_id)

    filename = import_file.name.lower()

    try:
        if filename.endswith('.csv'):
            rows = _parse_csv(import_file)
        elif filename.endswith(('.xlsx', '.xls')):
            rows = _parse_excel(import_file)
        else:
            messages.error(request, 'Unsupported file type. Please upload .csv, .xlsx, or .xls.')
            return redirect('coordinator:masterlist_by_section', section_id=section_id)

        # Create audit batch record
        batch = GradeUploadBatch.objects.create(
            section=section,
            grade_level=section.grade_level,
            school_year=section.school_year,
            program=section.program,
            quarter=quarter,
            uploaded_by=request.user,
            original_filename=import_file.name,
            status='processing',
            total_rows=len(rows),
        )

        saved, failed, error_details = _process_grade_rows(rows, section, quarter, batch)

        # Update batch with results
        batch.rows_saved = saved
        batch.rows_failed = failed
        batch.error_log = '\n'.join(error_details)
        if failed == 0:
            batch.status = 'done'
        elif saved > 0:
            batch.status = 'partial'
        else:
            batch.status = 'failed'
        batch.save(update_fields=['rows_saved', 'rows_failed', 'error_log', 'status'])

        if saved:
            messages.success(
                request,
                f'Successfully saved grades for {saved} student(s) — {"Final Grade" if quarter == 5 else f"Quarter {quarter}"}.'
            )
        if failed:
            messages.warning(
                request,
                f'{failed} row(s) had issues: {" | ".join(error_details[:5])}'
                + (' (and more…)' if len(error_details) > 5 else '')
            )
        if not saved and not failed:
            messages.info(request, 'The file was empty or had no processable rows.')

    except Exception as e:
        messages.error(request, f'Error reading file: {str(e)}')

    return redirect('coordinator:masterlist_by_section', section_id=section_id)


def _parse_csv(file):
    """Parse an uploaded CSV file into a list of dicts."""
    decoded = file.read().decode('utf-8-sig')
    reader = csv.DictReader(io.StringIO(decoded))
    return [row for row in reader]


def _parse_excel(file):
    """Parse an uploaded Excel file into a list of dicts."""
    wb = openpyxl.load_workbook(file, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []
    headers = [str(h).strip() if h is not None else '' for h in rows[0]]
    result = []
    for row in rows[1:]:
        if all(cell is None for cell in row):
            continue
        result.append(dict(zip(headers, row)))
    return result


def _process_grade_rows(rows, section, quarter, batch):
    """
    For each row in the uploaded file:
      - Match LRN to a Student
      - Match column headers to Subject records (by name or code, case-insensitive)
      - Upsert AcademicPerformance records

    Returns (students_saved, students_failed, error_list).
    """
    saved = 0
    failed = 0
    error_details = []

    # Build case-insensitive subject lookup: name -> Subject, code -> Subject
    subject_program = _resolve_subject_program(section)
    subjects = Subject.objects.filter(program=subject_program, is_active=True)
    subject_map = {}
    for subj in subjects:
        subject_map[subj.name.strip().lower()] = subj
        subject_map[subj.code.strip().lower()] = subj

    for i, row in enumerate(rows, start=2):  # row 1 is the header
        # Accept 'LRN', 'lrn', or 'student_lrn' as the LRN column
        lrn = str(
            row.get('LRN') or row.get('lrn') or row.get('student_lrn') or ''
        ).strip()

        if not lrn:
            failed += 1
            error_details.append(f'Row {i}: missing LRN — skipped')
            continue

        try:
            student = Student.objects.get(lrn=lrn)
        except Student.DoesNotExist:
            failed += 1
            error_details.append(f'Row {i}: LRN {lrn} not found in the system')
            continue

        row_grade_count = 0
        row_errors = []

        for col_name, raw_value in row.items():
            col_key = str(col_name).strip().lower()
            # Skip the LRN column itself
            if col_key in ('lrn', 'student_lrn') or not col_name:
                continue

            subject = subject_map.get(col_key)
            if subject is None:
                continue  # Unknown column — silently skip

            if raw_value in (None, ''):
                continue  # No grade entered — skip this cell

            try:
                grade = Decimal(str(raw_value)).quantize(Decimal('0.01'))
                if grade < 0 or grade > 100:
                    row_errors.append(
                        f'Row {i} [{subject.code}]: {grade} is out of range (0–100)'
                    )
                    continue
            except InvalidOperation:
                row_errors.append(
                    f'Row {i} [{subject.code}]: invalid value "{raw_value}"'
                )
                continue

            try:
                with transaction.atomic():
                    AcademicPerformance.objects.update_or_create(
                        student=student,
                        school_year=section.school_year,
                        grade_level=section.grade_level,
                        subject=subject,
                        quarter=quarter,
                        defaults={
                            'grade': grade,
                            'section': section,
                            'program': section.program,
                            'upload_batch': batch,
                        },
                    )
                row_grade_count += 1
            except Exception as e:
                row_errors.append(f'Row {i} [{subject.code}]: {e}')

        if row_errors:
            error_details.extend(row_errors)
            # Only count as failed if NO grades were saved for this student
            if row_grade_count == 0:
                failed += 1
            else:
                saved += 1  # partial save still counts the student
        elif row_grade_count > 0:
            saved += 1
        # If row_grade_count == 0 and no errors, the row had no grade columns — skip silently

    return saved, failed, error_details


# ─────────────────────────────────────────────────────────────
# EXPORT — Excel
# ─────────────────────────────────────────────────────────────

@coordinator_required
def export_masterlist_excel(request, section_id):
    """Export masterlist to Excel"""
    try:
        user_profile = request.user.profile
        program = user_profile.program

        section = get_object_or_404(
            Section.objects.select_related('program', 'school_year', 'adviser'),
            id=section_id,
            program=program
        )

        program_selections = ProgramSelection.objects.filter(
            assigned_section_id=section_id,
            admin_approved=True
        ).select_related(
            'student',
            'student__student_data',
            'student__academic_data'
        ).order_by('student__student_data__last_name', 'student__student_data__first_name')

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = f"{section.name} Masterlist"

        ws.merge_cells('A1:H1')
        ws['A1'] = f"{section.name} - Student Masterlist"
        ws['A1'].font = Font(bold=True, size=16)
        ws['A1'].alignment = Alignment(horizontal='center')

        ws.merge_cells('A2:H2')
        adviser_name = (
            f"{section.adviser.first_name} {section.adviser.last_name}"
            if section.adviser else 'Not assigned'
        )
        ws['A2'] = f"Program: {section.program.name} | Adviser: {adviser_name}"
        ws['A2'].alignment = Alignment(horizontal='center')

        headers = ['#', 'LRN', 'Last Name', 'First Name', 'Middle Name', 'Gender', 'Age', 'Overall Average']
        header_fill = PatternFill(start_color="991B1B", end_color="991B1B", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF")

        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=4, column=col_num)
            cell.value = header
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center')

        for idx, ps in enumerate(program_selections, 1):
            student = ps.student
            student_info = getattr(student, 'student_data', None)
            academic_info = getattr(student, 'academic_data', None)

            if student_info:
                row = idx + 4
                ws.cell(row=row, column=1, value=idx)
                ws.cell(row=row, column=2, value=student.lrn)
                ws.cell(row=row, column=3, value=student_info.last_name)
                ws.cell(row=row, column=4, value=student_info.first_name)
                ws.cell(row=row, column=5, value=student_info.middle_name or '')
                ws.cell(row=row, column=6, value=student_info.gender.capitalize())
                ws.cell(row=row, column=7, value=student_info.age)
                ws.cell(row=row, column=8, value=round(academic_info.overall_average, 2) if academic_info and academic_info.overall_average else 'N/A')

        col_widths = {'A': 5, 'B': 15, 'C': 20, 'D': 20, 'E': 20, 'F': 10, 'G': 8, 'H': 15}
        for col, width in col_widths.items():
            ws.column_dimensions[col].width = width

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="{section.name}_Masterlist.xlsx"'
        wb.save(response)
        return response
    except Exception as e:
        return HttpResponse(f"Error exporting: {str(e)}", status=500)


# ─────────────────────────────────────────────────────────────
# EXPORT — PDF
# ─────────────────────────────────────────────────────────────

@coordinator_required
def export_masterlist_pdf(request, section_id):
    """Export masterlist to PDF"""
    try:
        user_profile = request.user.profile
        program = user_profile.program

        section = get_object_or_404(
            Section.objects.select_related('program', 'school_year', 'adviser'),
            id=section_id,
            program=program
        )

        program_selections = ProgramSelection.objects.filter(
            assigned_section_id=section_id,
            admin_approved=True
        ).select_related(
            'student',
            'student__student_data',
            'student__academic_data'
        ).order_by('student__student_data__last_name', 'student__student_data__first_name')

        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{section.name}_Masterlist.pdf"'

        doc = SimpleDocTemplate(response, pagesize=landscape(letter))
        elements = []
        styles = getSampleStyleSheet()

        adviser_name = (
            f"{section.adviser.first_name} {section.adviser.last_name}"
            if section.adviser else 'Not assigned'
        )

        title = Paragraph(f"<b>{section.name} - Student Masterlist</b>", styles['Title'])
        elements.append(title)
        elements.append(Spacer(1, 0.2 * inch))

        info = Paragraph(f"Program: {section.program.name} | Adviser: {adviser_name}", styles['Normal'])
        elements.append(info)
        elements.append(Spacer(1, 0.3 * inch))

        data = [['#', 'LRN', 'Last Name', 'First Name', 'M.I.', 'Gender', 'Age', 'Average']]

        for idx, ps in enumerate(program_selections, 1):
            student = ps.student
            student_info = getattr(student, 'student_data', None)
            academic_info = getattr(student, 'academic_data', None)

            if student_info:
                data.append([
                    str(idx),
                    student.lrn,
                    student_info.last_name,
                    student_info.first_name,
                    student_info.middle_name[0] if student_info.middle_name else '',
                    student_info.gender.capitalize(),
                    str(student_info.age) if student_info.age else '',
                    f"{academic_info.overall_average:.2f}" if academic_info and academic_info.overall_average else 'N/A'
                ])

        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#991B1B')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
        ]))

        elements.append(table)
        doc.build(elements)
        return response
    except Exception as e:
        return HttpResponse(f"Error exporting PDF: {str(e)}", status=500)


# ─────────────────────────────────────────────────────────────
# EXPORT — Word (DOCX)
# ─────────────────────────────────────────────────────────────

@coordinator_required
def export_masterlist_docx(request, section_id):
    """Export masterlist to Word document"""
    try:
        user_profile = request.user.profile
        program = user_profile.program

        section = get_object_or_404(
            Section.objects.select_related('program', 'school_year', 'adviser'),
            id=section_id,
            program=program
        )

        program_selections = ProgramSelection.objects.filter(
            assigned_section_id=section_id,
            admin_approved=True
        ).select_related(
            'student',
            'student__student_data',
            'student__academic_data'
        ).order_by('student__student_data__last_name', 'student__student_data__first_name')

        doc = Document()

        title = doc.add_heading(f'{section.name} - Student Masterlist', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        adviser_name = (
            f"{section.adviser.first_name} {section.adviser.last_name}"
            if section.adviser else 'Not assigned'
        )
        info = doc.add_paragraph()
        info.add_run(f"Program: {section.program.name} | Adviser: {adviser_name}").bold = True
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = doc.add_table(rows=1, cols=8)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        headers = ['#', 'LRN', 'Last Name', 'First Name', 'Middle Name', 'Gender', 'Age', 'Average']
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)

        for idx, ps in enumerate(program_selections, 1):
            student = ps.student
            student_info = getattr(student, 'student_data', None)
            academic_info = getattr(student, 'academic_data', None)

            if student_info:
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = student.lrn
                row_cells[2].text = student_info.last_name
                row_cells[3].text = student_info.first_name
                row_cells[4].text = student_info.middle_name or ''
                row_cells[5].text = student_info.gender.capitalize()
                row_cells[6].text = str(student_info.age) if student_info.age else ''
                row_cells[7].text = (
                    f"{academic_info.overall_average:.2f}"
                    if academic_info and academic_info.overall_average else 'N/A'
                )

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{section.name}_Masterlist.docx"'
        doc.save(response)
        return response
    except Exception as e:
        return HttpResponse(f"Error exporting DOCX: {str(e)}", status=500)