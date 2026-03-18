import json
from django.shortcuts import render
from django.http import HttpResponse
from django.db.models import Q
from admin_app.decorators import coordinator_required
from admin_app.models import Section, SchoolYear
from enrollment_app.models import ProgramSelection
from coordinator_app.models import CoordinatorActivityLog

import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


PROGRAM_NAMES = {
    'STE': 'Science, Technology & Engineering',
    'REGULAR': 'General Academic Curriculum',
    'SPFL': 'Special Program in Foreign Language',
    'SPTVE': 'Special Program in Technical-Vocational Education',
    'OHSP': 'Online Hospitality & Service Program',
    'SNED': 'Special Needs Education Program',
}

SUBJECT_FIELDS = [
    'mathematics', 'science', 'english', 'filipino',
    'araling_panlipunan', 'edukasyon_sa_pagpapakatao',
    'edukasyon_pangkabuhayan', 'mapeh'
]

SUBJECT_LABELS = [
    'Math', 'Science', 'English', 'Filipino', 'AP', 'ESP', 'TLE', 'MAPEH'
]


def get_program_code(program_obj):
    """Extract the program code from various formats."""
    if program_obj is None:
        return 'STE'
    program_str = str(program_obj)
    if ' - ' in program_str:
        return program_str.split(' - ')[0].strip().upper()
    return program_str.strip().upper()


def _get_coordinator_context(request):
    """Get common coordinator context: program_obj, program_code, school_year."""
    program_obj = None
    if hasattr(request.user, 'profile') and hasattr(request.user.profile, 'program'):
        program_obj = request.user.profile.program
    program_code = get_program_code(program_obj)
    school_year = SchoolYear.objects.filter(is_active=True).first()
    return program_obj, program_code, school_year


def _log_report_activity(request, report_type, output_format, record_count=0, extra_info=None):
    """
    Helper to log report generation/download activity.
    """
    program_obj, program_code, school_year = _get_coordinator_context(request)
    
    # Determine action type
    if 'template' in report_type.lower() or 'form' in report_type.lower():
        action = 'template_downloaded'
    else:
        action = 'report_generated'
    
    description = f'Generated {report_type} ({output_format.upper()})'
    if record_count:
        description += f' with {record_count} records'
    if extra_info:
        description += f' - {extra_info}'
    
    CoordinatorActivityLog.log(
        user=request.user,
        action=action,
        description=description,
        category='report',
        program=program_obj,
        metadata={
            'report_type': report_type,
            'format': output_format,
            'record_count': record_count,
            'school_year': str(school_year) if school_year else None
        },
        ip_address=request.META.get('REMOTE_ADDR')
    )


def _get_base_selections(program_code, school_year, active_grade=None):
    filters = {'selected_program_code': program_code}
    if school_year:
        filters['school_year'] = school_year
    qs = ProgramSelection.objects.filter(**filters).select_related(
        'student', 'student__student_data', 'student__academic_data'
    ).order_by('student__student_data__last_name', 'student__student_data__first_name')
    if active_grade and active_grade != 'all':
       qs = qs.filter(assigned_section__grade_level__code=active_grade)
    return qs

def _calculate_gwa(academic_data):
    """Calculate GWA from AcademicData. Returns float or None."""
    grades = []
    for field in SUBJECT_FIELDS:
        val = getattr(academic_data, field, None)
        if val is not None:
            try:
                grades.append(float(val))
            except (ValueError, TypeError):
                pass
    return sum(grades) / len(grades) if grades else None


def _get_enrollment_status(ps):
    """Get human-readable enrollment status."""
    if ps.admin_approved:
        return 'Approved'
    elif ps.admin_rejected:
        return 'Rejected'
    elif hasattr(ps, 'student') and ps.student.enrollment_status == 'under_review':
        return 'Under Review'
    return 'Pending'


def _get_section_name(ps, sections_map):
    """Get section name from assigned_section id."""
    if ps.assigned_section:
        return sections_map.get(ps.assigned_section, ps.assigned_section)
    return 'Unassigned'


# ──────────────────────────────────────────────
# Main Reports Page
# ──────────────────────────────────────────────

@coordinator_required
def reports(request):
    program_obj, program_code, school_year = _get_coordinator_context(request)

    # Get base selections for this program
    active_grade = request.session.get('active_grade_code')
    base_selections = _get_base_selections(program_code, school_year, active_grade)
    
    # === STATISTICS COUNTS ===
    total_students = base_selections.count()
    approved_count = base_selections.filter(admin_approved=True).count()
    pending_count = base_selections.filter(admin_approved=False, admin_rejected=False).count()
    rejected_count = base_selections.filter(admin_rejected=True).count()
    
    # === SECTIONS DATA ===
    sections = []
    sections_list = []  # For JSON
    sections_count = 0
    if program_obj and school_year:
        active_grade = request.session.get('active_grade_code')
        section_filter = {'program': program_obj, 'school_year': school_year}
        if active_grade and active_grade != 'all':
            section_filter['grade_level__code'] = active_grade
        section_qs = Section.objects.filter(**section_filter).order_by('created_at')
        sections_count = section_qs.count()
        
        # Build sections with current_students count
        for sec in section_qs:
            actual_count = sec.get_actual_count()
            sections.append({
                'id': sec.id,
                'name': sec.name,
                'current_students': actual_count,
                'max_students': sec.max_students,
            })
            sections_list.append({
                'id': sec.id,
                'name': sec.name,
            })
    
    # === ACADEMIC COUNTS (Honor Roll & At-Risk) ===
    honor_count = 0
    at_risk_count = 0
    
    # Only count approved students for academic stats
    approved_selections = base_selections.filter(admin_approved=True)
    for ps in approved_selections:
        academic_data = getattr(ps.student, 'academic_data', None)
        if academic_data:
            gwa = _calculate_gwa(academic_data)
            if gwa:
                if gwa >= 90:
                    honor_count += 1
                elif gwa < 80:
                    at_risk_count += 1

    context = {
        'user': request.user,
        'program': program_code,
        'program_full_name': PROGRAM_NAMES.get(program_code, program_code),
        
        # Statistics for Quick Stats cards
        'total_students': total_students,
        'approved_count': approved_count,
        'pending_count': pending_count,
        'rejected_count': rejected_count,
        'sections_count': sections_count,
        
        # Sections for dropdowns (list of dicts with id, name, current_students)
        'sections': sections,
        'sections_json': json.dumps(sections_list, default=str),
        
        # Academic stats
        'honor_count': honor_count,
        'at_risk_count': at_risk_count,
        
        # School year label
        'school_year_label': school_year.year_label if school_year else 'N/A',
        'active_school_year': school_year, 
    }

    return render(request, 'coordinator_app/reports.html', context)


# ──────────────────────────────────────────────
# Enrollment Report (PDF)
# ──────────────────────────────────────────────

@coordinator_required
def generate_enrollment_report(request):
    """Generate enrollment report as PDF, Excel, or Word."""
    try:
        program_obj, program_code, school_year = _get_coordinator_context(request)
        status_filter = request.GET.get('status', 'all')
        output_format = request.GET.get('format', 'pdf').lower()
        year_label = school_year.year_label if school_year else 'N/A'

        active_grade = request.session.get('active_grade_code')
        selections = _get_base_selections(program_code, school_year, active_grade)

        # Apply status filter
        if status_filter == 'approved':
            selections = selections.filter(admin_approved=True)
        elif status_filter == 'rejected':
            selections = selections.filter(admin_rejected=True)
        elif status_filter == 'pending':
            selections = selections.filter(admin_approved=False, admin_rejected=False)

        # Build sections map for name lookup
        sections_map = {}
        if program_obj and school_year:
            for s in Section.objects.filter(program=program_obj, school_year=school_year):
                sections_map[str(s.id)] = s.name

        program_full = PROGRAM_NAMES.get(program_code, program_code)
        filter_label = status_filter.capitalize() if status_filter != 'all' else 'All Applicants'

        # Collect data for all formats
        student_rows = []
        for idx, ps in enumerate(selections, 1):
            student_info = getattr(ps.student, 'student_data', None)
            if student_info:
                student_rows.append({
                    'num': idx,
                    'lrn': ps.student.lrn,
                    'last_name': student_info.last_name,
                    'first_name': student_info.first_name,
                    'gender': (student_info.gender or '').capitalize(),
                    'status': _get_enrollment_status(ps),
                    'section': _get_section_name(ps, sections_map),
                })

        if output_format == 'word':
            # Generate Word Document
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml
            
            doc_file = Document()
            
            # Title
            title = doc_file.add_heading(f'{program_code} Enrollment Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Subtitle
            info_para = doc_file.add_paragraph()
            info_run = info_para.add_run(f"Program: {program_full} | School Year: {year_label} | Filter: {filter_label}")
            info_run.bold = True
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc_file.add_paragraph('')
            
            # Create table
            table = doc_file.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            headers = ['#', 'LRN', 'Last Name', 'First Name', 'Gender', 'Status', 'Section']
            for i, header in enumerate(headers):
                cell = header_cells[i]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="991B1B"/>')
                cell._tc.get_or_add_tcPr().append(shading_elm)
            
            # Data rows
            for row in student_rows:
                row_cells = table.add_row().cells
                row_cells[0].text = str(row['num'])
                row_cells[1].text = row['lrn']
                row_cells[2].text = row['last_name']
                row_cells[3].text = row['first_name']
                row_cells[4].text = row['gender']
                row_cells[5].text = row['status']
                row_cells[6].text = row['section']
            
            # Summary
            doc_file.add_paragraph('')
            summary = doc_file.add_paragraph()
            summary_run = summary.add_run(f"Total Records: {len(student_rows)}")
            summary_run.bold = True
            
            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{program_code}_Enrollment_Report.docx"'
            doc_file.save(response)
            
            # Log the report generation
            _log_report_activity(request, 'Enrollment Report', output_format, len(student_rows), f'Filter: {filter_label}')
            
            return response

        elif output_format == 'excel':
            # Generate Excel
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Enrollment Report"

            # Title
            ws.merge_cells('A1:G1')
            ws['A1'] = f"{program_code} Enrollment Report"
            ws['A1'].font = Font(bold=True, size=16)
            ws['A1'].alignment = Alignment(horizontal='center')

            ws.merge_cells('A2:G2')
            ws['A2'] = f"Program: {program_full} | School Year: {year_label} | Filter: {filter_label}"
            ws['A2'].alignment = Alignment(horizontal='center')

            # Headers
            headers = ['#', 'LRN', 'Last Name', 'First Name', 'Gender', 'Status', 'Section']
            header_fill = PatternFill(start_color="991B1B", end_color="991B1B", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF")
            for col_num, header in enumerate(headers, 1):
                cell = ws.cell(row=4, column=col_num)
                cell.value = header
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center')

            # Data
            row_num = 5
            for row in student_rows:
                ws.cell(row=row_num, column=1, value=row['num'])
                ws.cell(row=row_num, column=2, value=row['lrn'])
                ws.cell(row=row_num, column=3, value=row['last_name'])
                ws.cell(row=row_num, column=4, value=row['first_name'])
                ws.cell(row=row_num, column=5, value=row['gender'])
                ws.cell(row=row_num, column=6, value=row['status'])
                ws.cell(row=row_num, column=7, value=row['section'])
                row_num += 1

            # Column widths
            ws.column_dimensions['A'].width = 5
            ws.column_dimensions['B'].width = 15
            ws.column_dimensions['C'].width = 18
            ws.column_dimensions['D'].width = 18
            ws.column_dimensions['E'].width = 10
            ws.column_dimensions['F'].width = 12
            ws.column_dimensions['G'].width = 15

            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = f'attachment; filename="{program_code}_Enrollment_Report.xlsx"'
            wb.save(response)
            
            # Log the report generation
            _log_report_activity(request, 'Enrollment Report', output_format, len(student_rows), f'Filter: {filter_label}')
            
            return response

        else:
            # Generate PDF (default)
            response = HttpResponse(content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{program_code}_Enrollment_Report.pdf"'

            doc = SimpleDocTemplate(response, pagesize=landscape(letter))
            elements = []
            styles = getSampleStyleSheet()

            # Title
            title = Paragraph(
                f"<b>{program_code} Enrollment Report</b>",
                styles['Title']
            )
            elements.append(title)
            elements.append(Spacer(1, 0.1 * inch))

            # Subtitle
            info = Paragraph(
                f"Program: {program_full} | School Year: {year_label} | Filter: {filter_label}",
                styles['Normal']
            )
            elements.append(info)
            elements.append(Spacer(1, 0.3 * inch))

            # Table data
            data = [['#', 'LRN', 'Last Name', 'First Name', 'Gender', 'Status', 'Section']]

            for row in student_rows:
                data.append([
                    str(row['num']),
                    row['lrn'],
                    row['last_name'],
                    row['first_name'],
                    row['gender'],
                    row['status'],
                    row['section'],
                ])

            if len(data) == 1:
                data.append(['', '', '', 'No records found', '', '', ''])

            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#991B1B')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
            ]))

            elements.append(table)

            # Summary footer
            elements.append(Spacer(1, 0.3 * inch))
            elements.append(Paragraph(f"<b>Total Records: {len(data) - 1}</b>", styles['Normal']))

            doc.build(elements)
            
            # Log the report generation
            _log_report_activity(request, 'Enrollment Report', output_format, len(student_rows), f'Filter: {filter_label}')
            
            return response

    except Exception as e:
        return HttpResponse(f"Error generating enrollment report: {str(e)}", status=500)


# ──────────────────────────────────────────────
# Academic Performance Report (PDF/Excel)
# ──────────────────────────────────────────────

@coordinator_required
def generate_academic_report(request):
    """Generate academic performance report as PDF or Excel."""
    try:
        program_obj, program_code, school_year = _get_coordinator_context(request)
        # Accept both 'filter' and 'gwa_filter' parameters
        gwa_filter = request.GET.get('filter', request.GET.get('gwa_filter', 'all'))
        output_format = request.GET.get('format', 'excel').lower()
        year_label = school_year.year_label if school_year else 'N/A'

        active_grade = request.session.get('active_grade_code')
        selections = _get_base_selections(program_code, school_year, active_grade).filter(admin_approved=True)
        program_full = PROGRAM_NAMES.get(program_code, program_code)
        
        # Map filter values
        filter_map = {
            'all': ('All Students', None, None),
            'honors': ('Honor Students (GWA ≥90)', 90, None),
            'above90': ('GWA 90 & Above', 90, None),
            'top20': ('Top 20 Students', None, None),
            'below85': ('GWA Below 85', None, 85),
            'below80': ('At-Risk (GWA <80)', None, 80),
        }
        filter_label, min_gwa, max_gwa = filter_map.get(gwa_filter, ('All Students', None, None))

        # Build student data list with GWA for filtering and sorting
        student_data = []
        for ps in selections:
            student_info = getattr(ps.student, 'student_data', None)
            academic_info = getattr(ps.student, 'academic_data', None)
            if not student_info or not academic_info:
                continue
            gwa = _calculate_gwa(academic_info)
            if gwa is None:
                continue
            # Apply GWA filter
            if min_gwa and gwa < min_gwa:
                continue
            if max_gwa and gwa >= max_gwa:
                continue
            student_data.append({
                'ps': ps,
                'student_info': student_info,
                'academic_info': academic_info,
                'gwa': gwa,
            })

        # Sort by GWA descending
        student_data.sort(key=lambda x: x['gwa'], reverse=True)

        # Apply top20 limit
        if gwa_filter == 'top20':
            student_data = student_data[:20]

        if output_format == 'pdf':
            # Generate PDF
            response = HttpResponse(content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{program_code}_Academic_Rankings.pdf"'

            doc = SimpleDocTemplate(response, pagesize=landscape(letter))
            elements = []
            styles = getSampleStyleSheet()

            elements.append(Paragraph(f"<b>{program_code} Academic Rankings</b>", styles['Title']))
            elements.append(Spacer(1, 0.1 * inch))
            elements.append(Paragraph(f"Program: {program_full} | School Year: {year_label} | Filter: {filter_label}", styles['Normal']))
            elements.append(Spacer(1, 0.3 * inch))

            # Table
            data = [['Rank', 'LRN', 'Last Name', 'First Name', 'GWA']]
            for rank, item in enumerate(student_data, 1):
                data.append([
                    str(rank),
                    item['ps'].student.lrn,
                    item['student_info'].last_name,
                    item['student_info'].first_name,
                    f"{item['gwa']:.2f}",
                ])

            if len(data) == 1:
                data.append(['', '', '', 'No records found', ''])

            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#991B1B')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
            ]))
            elements.append(table)
            elements.append(Spacer(1, 0.3 * inch))
            elements.append(Paragraph(f"<b>Total Records: {len(data) - 1}</b>", styles['Normal']))

            doc.build(elements)
            
            # Log the report generation
            _log_report_activity(request, 'Academic Report', output_format, len(student_data), f'Filter: {filter_label}')
            
            return response

        else:
            # Generate Excel (default)
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = f"{program_code} Academic Performance"

            header_cols = 4 + len(SUBJECT_FIELDS) + 1
            ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=header_cols)
            ws['A1'] = f"{program_code} Academic Performance Report"
            ws['A1'].font = Font(bold=True, size=16)
            ws['A1'].alignment = Alignment(horizontal='center')

            ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=header_cols)
            ws['A2'] = f"Program: {program_full} | School Year: {year_label} | Filter: {filter_label}"
            ws['A2'].alignment = Alignment(horizontal='center')

            headers = ['Rank', 'LRN', 'Last Name', 'First Name'] + SUBJECT_LABELS + ['Overall Avg']
            header_fill = PatternFill(start_color="991B1B", end_color="991B1B", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF")

            for col_num, header in enumerate(headers, 1):
                cell = ws.cell(row=4, column=col_num)
                cell.value = header
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center')

            row_num = 5
            for rank, item in enumerate(student_data, 1):
                ws.cell(row=row_num, column=1, value=rank)
                ws.cell(row=row_num, column=2, value=item['ps'].student.lrn)
                ws.cell(row=row_num, column=3, value=item['student_info'].last_name)
                ws.cell(row=row_num, column=4, value=item['student_info'].first_name)

                for i, field in enumerate(SUBJECT_FIELDS):
                    val = getattr(item['academic_info'], field, None)
                    ws.cell(row=row_num, column=5 + i, value=round(float(val), 2) if val is not None else 'N/A')

                ws.cell(row=row_num, column=5 + len(SUBJECT_FIELDS), value=round(item['gwa'], 2))
                row_num += 1

            ws.cell(row=row_num + 1, column=1, value=f"Total Records: {len(student_data)}")
            ws.cell(row=row_num + 1, column=1).font = Font(bold=True)

            ws.column_dimensions['A'].width = 6
            ws.column_dimensions['B'].width = 15
            ws.column_dimensions['C'].width = 18
            ws.column_dimensions['D'].width = 18
            for i in range(len(SUBJECT_FIELDS)):
                col_letter = openpyxl.utils.get_column_letter(5 + i)
                ws.column_dimensions[col_letter].width = 10
            ws.column_dimensions[openpyxl.utils.get_column_letter(5 + len(SUBJECT_FIELDS))].width = 12

            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = f'attachment; filename="{program_code}_Academic_Performance.xlsx"'
            wb.save(response)
            
            # Log the report generation
            _log_report_activity(request, 'Academic Report', output_format, len(student_data), f'Filter: {filter_label}')
            
            return response

    except Exception as e:
        return HttpResponse(f"Error generating academic report: {str(e)}", status=500)


# ──────────────────────────────────────────────
# Section Assignment Report (PDF/Excel)
# ──────────────────────────────────────────────

@coordinator_required
def generate_section_report(request):
    """Generate section assignment report as PDF or Excel."""
    try:
        program_obj, program_code, school_year = _get_coordinator_context(request)
        section_filter = request.GET.get('section', request.GET.get('section_id', 'all'))
        output_format = request.GET.get('format', 'pdf').lower()
        year_label = school_year.year_label if school_year else 'N/A'
        program_full = PROGRAM_NAMES.get(program_code, program_code)

        # Get sections
        section_qs = Section.objects.none()
        if program_obj and school_year:
            section_qs = Section.objects.filter(
                program=program_obj, school_year=school_year
            ).order_by('created_at')

        if section_filter != 'all':
            section_qs = section_qs.filter(id=section_filter)

        # Collect all data for both formats
        all_section_data = []
        total_students = 0

        for section in section_qs:
            # Filter by section FK
            selections = ProgramSelection.objects.filter(
                assigned_section=section,
                admin_approved=True,
                selected_program_code__icontains=program_code
            ).select_related(
                'student', 'student__student_data', 'student__academic_data'
            ).order_by('student__student_data__last_name', 'student__student_data__first_name')

            students = []
            for idx, ps in enumerate(selections, 1):
                student_info = getattr(ps.student, 'student_data', None)
                academic_info = getattr(ps.student, 'academic_data', None)
                if student_info:
                    full_name = f"{student_info.last_name}, {student_info.first_name}"
                    if student_info.middle_name:
                        full_name += f" {student_info.middle_name[0]}."
                    gwa = None
                    if academic_info and academic_info.overall_average:
                        gwa = float(academic_info.overall_average)
                    elif academic_info:
                        gwa = _calculate_gwa(academic_info)
                    students.append({
                        'num': idx,
                        'lrn': ps.student.lrn,
                        'name': full_name,
                        'gender': (student_info.gender or '').capitalize(),
                        'gwa': gwa,
                    })
                    total_students += 1

            all_section_data.append({
                'name': section.name,
                'capacity': f"{section.get_actual_count()}/{section.max_students}",
                'students': students,
            })

        if output_format == 'excel':
            # Generate Excel
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Class List"

            ws.merge_cells('A1:E1')
            ws['A1'] = f"{program_code} Class List by Section"
            ws['A1'].font = Font(bold=True, size=16)
            ws['A1'].alignment = Alignment(horizontal='center')

            ws.merge_cells('A2:E2')
            ws['A2'] = f"Program: {program_full} | School Year: {year_label}"
            ws['A2'].alignment = Alignment(horizontal='center')

            header_fill = PatternFill(start_color="991B1B", end_color="991B1B", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF")
            section_fill = PatternFill(start_color="FEE2E2", end_color="FEE2E2", fill_type="solid")

            row_num = 4
            for sec_data in all_section_data:
                # Section header
                ws.merge_cells(start_row=row_num, start_column=1, end_row=row_num, end_column=5)
                ws.cell(row=row_num, column=1, value=f"{sec_data['name']} ({sec_data['capacity']})")
                ws.cell(row=row_num, column=1).font = Font(bold=True, size=12)
                ws.cell(row=row_num, column=1).fill = section_fill
                row_num += 1

                # Column headers
                for col_num, header in enumerate(['#', 'LRN', 'Full Name', 'Gender', 'GWA'], 1):
                    cell = ws.cell(row=row_num, column=col_num)
                    cell.value = header
                    cell.fill = header_fill
                    cell.font = header_font
                    cell.alignment = Alignment(horizontal='center')
                row_num += 1

                if not sec_data['students']:
                    ws.cell(row=row_num, column=1, value="No students assigned")
                    row_num += 2
                    continue

                for student in sec_data['students']:
                    ws.cell(row=row_num, column=1, value=student['num'])
                    ws.cell(row=row_num, column=2, value=student['lrn'])
                    ws.cell(row=row_num, column=3, value=student['name'])
                    ws.cell(row=row_num, column=4, value=student['gender'])
                    ws.cell(row=row_num, column=5, value=f"{student['gwa']:.2f}" if student['gwa'] else 'N/A')
                    row_num += 1

                row_num += 1  # Spacer

            ws.cell(row=row_num, column=1, value=f"Total Students: {total_students}")
            ws.cell(row=row_num, column=1).font = Font(bold=True)

            ws.column_dimensions['A'].width = 5
            ws.column_dimensions['B'].width = 15
            ws.column_dimensions['C'].width = 30
            ws.column_dimensions['D'].width = 10
            ws.column_dimensions['E'].width = 10

            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = f'attachment; filename="{program_code}_Class_List.xlsx"'
            wb.save(response)
            
            # Log the report generation
            _log_report_activity(request, 'Section Report', output_format, total_students, f'Sections: {len(all_section_data)}')
            
            return response

        elif output_format == 'word':
            # Generate Word Document
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml
            
            doc_file = Document()
            
            # Title
            title = doc_file.add_heading(f'{program_code} Class List by Section', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Subtitle
            info_para = doc_file.add_paragraph()
            info_run = info_para.add_run(f"Program: {program_full} | School Year: {year_label}")
            info_run.bold = True
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc_file.add_paragraph('')
            
            for sec_data in all_section_data:
                # Section heading
                doc_file.add_heading(f"{sec_data['name']} ({sec_data['capacity']})", level=1)
                
                if not sec_data['students']:
                    doc_file.add_paragraph('No students assigned to this section.')
                    doc_file.add_paragraph('')
                    continue
                
                # Create table
                table = doc_file.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                
                # Header row
                header_cells = table.rows[0].cells
                headers = ['#', 'LRN', 'Full Name', 'Gender', 'GWA']
                for i, header in enumerate(headers):
                    cell = header_cells[i]
                    cell.text = header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)
                    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="991B1B"/>')
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                
                # Data rows
                for student in sec_data['students']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(student['num'])
                    row_cells[1].text = student['lrn']
                    row_cells[2].text = student['name']
                    row_cells[3].text = student['gender']
                    row_cells[4].text = f"{student['gwa']:.2f}" if student['gwa'] else 'N/A'
                
                doc_file.add_paragraph('')
            
            # Summary
            summary = doc_file.add_paragraph()
            summary_run = summary.add_run(f"Total Sections: {len(all_section_data)} | Total Students: {total_students}")
            summary_run.bold = True
            
            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{program_code}_Class_List.docx"'
            doc_file.save(response)
            
            # Log the report generation
            _log_report_activity(request, 'Section Report', output_format, total_students, f'Sections: {len(all_section_data)}')
            
            return response

        else:
            # Generate PDF (default)
            response = HttpResponse(content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{program_code}_Class_List.pdf"'

            doc = SimpleDocTemplate(response, pagesize=letter)
            elements = []
            styles = getSampleStyleSheet()

            elements.append(Paragraph(f"<b>{program_code} Class List by Section</b>", styles['Title']))
            elements.append(Spacer(1, 0.1 * inch))
            elements.append(Paragraph(f"Program: {program_full} | School Year: {year_label}", styles['Normal']))
            elements.append(Spacer(1, 0.3 * inch))

            for sec_data in all_section_data:
                # Section heading
                elements.append(Paragraph(f"<b>{sec_data['name']}</b> ({sec_data['capacity']})", styles['Heading2']))
                elements.append(Spacer(1, 0.1 * inch))

                if not sec_data['students']:
                    elements.append(Paragraph("No students assigned to this section.", styles['Normal']))
                    elements.append(Spacer(1, 0.2 * inch))
                    continue

                data = [['#', 'LRN', 'Full Name', 'Gender', 'GWA']]
                for student in sec_data['students']:
                    data.append([
                        str(student['num']),
                        student['lrn'],
                        student['name'],
                        student['gender'],
                        f"{student['gwa']:.2f}" if student['gwa'] else 'N/A',
                    ])

                table = Table(data, colWidths=[30, 100, 180, 60, 60])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#991B1B')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('ALIGN', (2, 1), (2, -1), 'LEFT'),
                ]))
                elements.append(table)
                elements.append(Spacer(1, 0.3 * inch))

            elements.append(Paragraph(f"<b>Total Sections: {len(all_section_data)} | Total Students: {total_students}</b>", styles['Normal']))

            doc.build(elements)
            
            # Log the report generation
            _log_report_activity(request, 'Section Report', output_format, total_students, f'Sections: {len(all_section_data)}')
            
            return response

    except Exception as e:
        return HttpResponse(f"Error generating section report: {str(e)}", status=500)


@coordinator_required
def generate_enrollment_analytics_report(request):
    """
    Generate Enrollment Analytics PDF with statistics, breakdowns, and summary metrics.
    """
    try:
        program_obj, program_code, school_year = _get_coordinator_context(request)
        program_full = PROGRAM_NAMES.get(program_code, program_code)
        year_label = str(school_year) if school_year else 'All Years'

        # Get all enrollments for this program
        active_grade = request.session.get('active_grade_code')
        selections = _get_base_selections(program_code, school_year, active_grade)

        # Calculate statistics
        total_students = selections.count()
        approved_count = selections.filter(admin_approved=True).count()
        pending_count = selections.filter(admin_approved=False, admin_rejected=False).count()
        rejected_count = selections.filter(admin_rejected=True).count()

        # Gender breakdown
        male_count = 0
        female_count = 0
        for ps in selections:
            if hasattr(ps, 'student') and hasattr(ps.student, 'student_data'):
                gender = getattr(ps.student.student_data, 'gender', '')
                if gender and gender.lower() == 'male':
                    male_count += 1
                elif gender and gender.lower() == 'female':
                    female_count += 1

        # GWA statistics
        gwa_list = []
        honors_count = 0
        at_risk_count = 0
        for ps in selections:
            if hasattr(ps, 'student') and hasattr(ps.student, 'academic_data'):
                gwa = _calculate_gwa(ps.student.academic_data)
                if gwa:
                    gwa_list.append(gwa)
                    if gwa <= 85:
                        honors_count += 1
                    if gwa > 80:
                        at_risk_count += 1

        avg_gwa = sum(gwa_list) / len(gwa_list) if gwa_list else 0
        min_gwa = min(gwa_list) if gwa_list else 0
        max_gwa = max(gwa_list) if gwa_list else 0

        # Section distribution
        sections = Section.objects.filter(program=program_obj)
        sections_map = {str(s.id): s.name for s in sections}
        section_counts = {}
        unassigned_count = 0
        for ps in selections:
            if ps.assigned_section:
                sec_name = sections_map.get(ps.assigned_section, ps.assigned_section)
                section_counts[sec_name] = section_counts.get(sec_name, 0) + 1
            else:
                unassigned_count += 1

        # Generate PDF
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{program_code}_Enrollment_Analytics.pdf"'

        doc = SimpleDocTemplate(response, pagesize=letter)
        elements = []
        styles = getSampleStyleSheet()

        # Title
        elements.append(Paragraph(f"<b>{program_code} Enrollment Analytics Report</b>", styles['Title']))
        elements.append(Spacer(1, 0.1 * inch))
        elements.append(Paragraph(f"Program: {program_full} | School Year: {year_label}", styles['Normal']))
        elements.append(Spacer(1, 0.3 * inch))

        # Summary Statistics
        elements.append(Paragraph("<b>Summary Statistics</b>", styles['Heading2']))
        elements.append(Spacer(1, 0.1 * inch))

        summary_data = [
            ['Metric', 'Value'],
            ['Total Enrollees', str(total_students)],
            ['Approved', f"{approved_count} ({(approved_count/total_students*100):.1f}%)" if total_students else '0'],
            ['Pending', f"{pending_count} ({(pending_count/total_students*100):.1f}%)" if total_students else '0'],
            ['Rejected', f"{rejected_count} ({(rejected_count/total_students*100):.1f}%)" if total_students else '0'],
        ]

        summary_table = Table(summary_data, colWidths=[200, 200])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#991B1B')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#FEF2F2')),
        ]))
        elements.append(summary_table)
        elements.append(Spacer(1, 0.3 * inch))

        # Gender Distribution
        elements.append(Paragraph("<b>Gender Distribution</b>", styles['Heading2']))
        elements.append(Spacer(1, 0.1 * inch))

        gender_data = [
            ['Gender', 'Count', 'Percentage'],
            ['Male', str(male_count), f"{(male_count/total_students*100):.1f}%" if total_students else '0%'],
            ['Female', str(female_count), f"{(female_count/total_students*100):.1f}%" if total_students else '0%'],
        ]

        gender_table = Table(gender_data, colWidths=[150, 100, 100])
        gender_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E40AF')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#EFF6FF')),
        ]))
        elements.append(gender_table)
        elements.append(Spacer(1, 0.3 * inch))

        # Academic Performance
        elements.append(Paragraph("<b>Academic Performance</b>", styles['Heading2']))
        elements.append(Spacer(1, 0.1 * inch))

        academic_data = [
            ['Metric', 'Value'],
            ['Average GWA', f"{avg_gwa:.2f}" if avg_gwa else 'N/A'],
            ['Highest GWA', f"{min_gwa:.2f}" if min_gwa else 'N/A'],
            ['Lowest GWA', f"{max_gwa:.2f}" if max_gwa else 'N/A'],
            ['Honor Students (≤85 GWA)', str(honors_count)],
            ['At-Risk Students (>80 GWA)', str(at_risk_count)],
        ]

        academic_table = Table(academic_data, colWidths=[200, 200])
        academic_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#065F46')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#ECFDF5')),
        ]))
        elements.append(academic_table)
        elements.append(Spacer(1, 0.3 * inch))

        # Section Distribution
        elements.append(Paragraph("<b>Section Distribution</b>", styles['Heading2']))
        elements.append(Spacer(1, 0.1 * inch))

        section_data = [['Section', 'Enrolled', 'Percentage']]
        for sec_name, count in sorted(section_counts.items()):
            pct = (count / total_students * 100) if total_students else 0
            section_data.append([sec_name, str(count), f"{pct:.1f}%"])
        if unassigned_count:
            pct = (unassigned_count / total_students * 100) if total_students else 0
            section_data.append(['Unassigned', str(unassigned_count), f"{pct:.1f}%"])

        if len(section_data) > 1:
            section_table = Table(section_data, colWidths=[150, 100, 100])
            section_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#7C3AED')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F5F3FF')),
            ]))
            elements.append(section_table)
        else:
            elements.append(Paragraph("No section assignments yet.", styles['Normal']))

        elements.append(Spacer(1, 0.5 * inch))
        elements.append(Paragraph(f"<i>Report generated on {__import__('datetime').datetime.now().strftime('%B %d, %Y at %I:%M %p')}</i>", styles['Normal']))

        doc.build(elements)
        
        # Log the report generation
        _log_report_activity(request, 'Enrollment Analytics Report', 'pdf', total_students, 'Enrollment Analytics')
        
        return response

    except Exception as e:
        return HttpResponse(f"Error generating enrollment analytics: {str(e)}", status=500)


@coordinator_required
def generate_section_analytics_report(request):
    """
    Generate Section Analytics PDF with capacity utilization, distribution, and comparison.
    """
    try:
        program_obj, program_code, school_year = _get_coordinator_context(request)
        program_full = PROGRAM_NAMES.get(program_code, program_code)
        year_label = str(school_year) if school_year else 'All Years'

        # Get sections for this program
        sections = Section.objects.filter(program=program_obj).order_by('name')
        active_grade = request.session.get('active_grade_code')
        selections = _get_base_selections(program_code, school_year, active_grade).filter(admin_approved=True)
        # Build section statistics
        sections_map = {str(s.id): s.name for s in sections}
        section_stats = []
        total_capacity = 0
        total_enrolled = 0

        for section in sections:
            capacity = section.max_students or 0
            enrolled = selections.filter(
                assigned_section=section
            ).count()

            total_capacity += capacity
            total_enrolled += enrolled

            utilization = (enrolled / capacity * 100) if capacity else 0
            section_stats.append({
                'name': section.name,
                'capacity': capacity,
                'enrolled': enrolled,
                'available': max(0, capacity - enrolled),
                'utilization': utilization,
            })

        # GWA by section
        section_gwa = {}
        for section in sections:
            sec_selections = selections.filter(
                assigned_section=section
            )
            gwa_list = []
            for ps in sec_selections:
                if hasattr(ps, 'student') and hasattr(ps.student, 'academic_data'):
                    gwa = _calculate_gwa(ps.student.academic_data)
                    if gwa:
                        gwa_list.append(gwa)
            section_gwa[section.name] = sum(gwa_list) / len(gwa_list) if gwa_list else None

        # Gender by section
        section_gender = {}
        for section in sections:
            sec_selections = selections.filter(
                assigned_section=section
            )
            male = 0
            female = 0
            for ps in sec_selections:
                if hasattr(ps, 'student') and hasattr(ps.student, 'student_data'):
                    gender = getattr(ps.student.student_data, 'gender', '')
                    if gender and gender.lower() == 'male':
                        male += 1
                    elif gender and gender.lower() == 'female':
                        female += 1
            section_gender[section.name] = {'male': male, 'female': female}

        # Generate PDF
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{program_code}_Section_Analytics.pdf"'

        doc = SimpleDocTemplate(response, pagesize=letter)
        elements = []
        styles = getSampleStyleSheet()

        # Title
        elements.append(Paragraph(f"<b>{program_code} Section Analytics Report</b>", styles['Title']))
        elements.append(Spacer(1, 0.1 * inch))
        elements.append(Paragraph(f"Program: {program_full} | School Year: {year_label}", styles['Normal']))
        elements.append(Spacer(1, 0.3 * inch))

        # Overall Summary
        elements.append(Paragraph("<b>Overall Summary</b>", styles['Heading2']))
        elements.append(Spacer(1, 0.1 * inch))

        overall_util = (total_enrolled / total_capacity * 100) if total_capacity else 0
        summary_data = [
            ['Metric', 'Value'],
            ['Total Sections', str(len(sections))],
            ['Total Capacity', str(total_capacity)],
            ['Total Enrolled', str(total_enrolled)],
            ['Available Slots', str(max(0, total_capacity - total_enrolled))],
            ['Overall Utilization', f"{overall_util:.1f}%"],
        ]

        summary_table = Table(summary_data, colWidths=[200, 200])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#991B1B')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#FEF2F2')),
        ]))
        elements.append(summary_table)
        elements.append(Spacer(1, 0.3 * inch))

        # Section Capacity Table
        elements.append(Paragraph("<b>Section Capacity & Utilization</b>", styles['Heading2']))
        elements.append(Spacer(1, 0.1 * inch))

        capacity_data = [['Section', 'Capacity', 'Enrolled', 'Available', 'Utilization']]
        for stat in section_stats:
            capacity_data.append([
                stat['name'],
                str(stat['capacity']),
                str(stat['enrolled']),
                str(stat['available']),
                f"{stat['utilization']:.1f}%"
            ])

        capacity_table = Table(capacity_data, colWidths=[100, 80, 80, 80, 80])
        capacity_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E40AF')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#EFF6FF')),
        ]))
        elements.append(capacity_table)
        elements.append(Spacer(1, 0.3 * inch))

        # Average GWA by Section
        elements.append(Paragraph("<b>Average GWA by Section</b>", styles['Heading2']))
        elements.append(Spacer(1, 0.1 * inch))

        gwa_data = [['Section', 'Average GWA']]
        for sec_name, avg_gwa in section_gwa.items():
            gwa_data.append([sec_name, f"{avg_gwa:.2f}" if avg_gwa else 'N/A'])

        gwa_table = Table(gwa_data, colWidths=[200, 150])
        gwa_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#065F46')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#ECFDF5')),
        ]))
        elements.append(gwa_table)
        elements.append(Spacer(1, 0.3 * inch))

        # Gender Distribution by Section
        elements.append(Paragraph("<b>Gender Distribution by Section</b>", styles['Heading2']))
        elements.append(Spacer(1, 0.1 * inch))

        gender_data = [['Section', 'Male', 'Female', 'Total']]
        for sec_name, counts in section_gender.items():
            total = counts['male'] + counts['female']
            gender_data.append([sec_name, str(counts['male']), str(counts['female']), str(total)])

        gender_table = Table(gender_data, colWidths=[120, 80, 80, 80])
        gender_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#7C3AED')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F5F3FF')),
        ]))
        elements.append(gender_table)

        elements.append(Spacer(1, 0.5 * inch))
        elements.append(Paragraph(f"<i>Report generated on {__import__('datetime').datetime.now().strftime('%B %d, %Y at %I:%M %p')}</i>", styles['Normal']))

        doc.build(elements)
        
        # Log the report generation
        _log_report_activity(request, 'Section Analytics Report', 'pdf', total_enrolled, f'Sections: {len(sections)}')
        
        return response

    except Exception as e:
        return HttpResponse(f"Error generating section analytics: {str(e)}", status=500)


# ──────────────────────────────────────────────
# Administrative Reports
# ──────────────────────────────────────────────

@coordinator_required
def generate_program_summary_report(request):
    """
    Generate Program Summary PDF with comprehensive overview of the program.
    """
    import datetime
    try:
        program_obj, program_code, school_year = _get_coordinator_context(request)
        program_full = PROGRAM_NAMES.get(program_code, program_code)
        year_label = school_year.year_label if school_year else 'N/A'
        period = request.GET.get('period', 'current')

        period_labels = {
            'current': 'Current Period',
            'q1': '1st Quarter',
            'q2': '2nd Quarter',
            'semester': 'Semester Summary',
            'annual': 'Annual Summary',
        }
        period_label = period_labels.get(period, 'Current Period')

        # Get all enrollments for this program
        active_grade = request.session.get('active_grade_code')
        selections = _get_base_selections(program_code, school_year, active_grade)

        # Statistics
        total_students = selections.count()
        approved_count = selections.filter(admin_approved=True).count()
        pending_count = selections.filter(admin_approved=False, admin_rejected=False).count()
        rejected_count = selections.filter(admin_rejected=True).count()

        # Gender breakdown
        male_count = 0
        female_count = 0
        for ps in selections:
            if hasattr(ps, 'student') and hasattr(ps.student, 'student_data'):
                gender = getattr(ps.student.student_data, 'gender', '')
                if gender and gender.lower() == 'male':
                    male_count += 1
                elif gender and gender.lower() == 'female':
                    female_count += 1

        # GWA statistics
        gwa_list = []
        honors_count = 0
        at_risk_count = 0
        for ps in selections.filter(admin_approved=True):
            if hasattr(ps, 'student') and hasattr(ps.student, 'academic_data'):
                gwa = _calculate_gwa(ps.student.academic_data)
                if gwa:
                    gwa_list.append(gwa)
                    if gwa >= 90:
                        honors_count += 1
                    if gwa < 80:
                        at_risk_count += 1

        avg_gwa = sum(gwa_list) / len(gwa_list) if gwa_list else 0

        # Sections
        sections = Section.objects.filter(program=program_obj, school_year=school_year) if program_obj and school_year else []
        total_capacity = sum(s.max_students or 0 for s in sections)
        sections_count = len(sections)

        # Generate PDF
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{program_code}_Program_Summary.pdf"'

        doc = SimpleDocTemplate(response, pagesize=letter)
        elements = []
        styles = getSampleStyleSheet()

        # Title
        elements.append(Paragraph(f"<b>{program_code} Program Summary Report</b>", styles['Title']))
        elements.append(Spacer(1, 0.1 * inch))
        elements.append(Paragraph(f"Program: {program_full}", styles['Normal']))
        elements.append(Paragraph(f"School Year: {year_label} | Period: {period_label}", styles['Normal']))
        elements.append(Spacer(1, 0.3 * inch))

        # Program Overview
        elements.append(Paragraph("<b>Program Overview</b>", styles['Heading2']))
        elements.append(Spacer(1, 0.1 * inch))

        overview_data = [
            ['Metric', 'Value'],
            ['Program Name', program_full],
            ['Program Code', program_code],
            ['School Year', year_label],
            ['Total Sections', str(sections_count)],
            ['Total Capacity', str(total_capacity)],
        ]

        overview_table = Table(overview_data, colWidths=[200, 250])
        overview_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#991B1B')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#FEF2F2')),
        ]))
        elements.append(overview_table)
        elements.append(Spacer(1, 0.3 * inch))

        # Enrollment Statistics
        elements.append(Paragraph("<b>Enrollment Statistics</b>", styles['Heading2']))
        elements.append(Spacer(1, 0.1 * inch))

        enrollment_data = [
            ['Status', 'Count', 'Percentage'],
            ['Total Applicants', str(total_students), '100%'],
            ['Approved', str(approved_count), f"{(approved_count/total_students*100):.1f}%" if total_students else '0%'],
            ['Pending', str(pending_count), f"{(pending_count/total_students*100):.1f}%" if total_students else '0%'],
            ['Rejected', str(rejected_count), f"{(rejected_count/total_students*100):.1f}%" if total_students else '0%'],
        ]

        enrollment_table = Table(enrollment_data, colWidths=[150, 100, 100])
        enrollment_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E40AF')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#EFF6FF')),
        ]))
        elements.append(enrollment_table)
        elements.append(Spacer(1, 0.3 * inch))

        # Demographics
        elements.append(Paragraph("<b>Student Demographics</b>", styles['Heading2']))
        elements.append(Spacer(1, 0.1 * inch))

        demo_data = [
            ['Category', 'Count', 'Percentage'],
            ['Male Students', str(male_count), f"{(male_count/total_students*100):.1f}%" if total_students else '0%'],
            ['Female Students', str(female_count), f"{(female_count/total_students*100):.1f}%" if total_students else '0%'],
        ]

        demo_table = Table(demo_data, colWidths=[150, 100, 100])
        demo_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#065F46')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#ECFDF5')),
        ]))
        elements.append(demo_table)
        elements.append(Spacer(1, 0.3 * inch))

        # Academic Performance
        elements.append(Paragraph("<b>Academic Performance Summary</b>", styles['Heading2']))
        elements.append(Spacer(1, 0.1 * inch))

        academic_data = [
            ['Metric', 'Value'],
            ['Average GWA', f"{avg_gwa:.2f}" if avg_gwa else 'N/A'],
            ['Honor Students (GWA ≥90)', str(honors_count)],
            ['At-Risk Students (GWA <80)', str(at_risk_count)],
            ['Students with GWA Data', str(len(gwa_list))],
        ]

        academic_table = Table(academic_data, colWidths=[200, 150])
        academic_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#7C3AED')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F5F3FF')),
        ]))
        elements.append(academic_table)

        elements.append(Spacer(1, 0.5 * inch))
        elements.append(Paragraph(f"<i>Report generated on {datetime.datetime.now().strftime('%B %d, %Y at %I:%M %p')}</i>", styles['Normal']))
        elements.append(Paragraph(f"<i>Generated by: {request.user.get_full_name() or request.user.username}</i>", styles['Normal']))

        doc.build(elements)
        
        # Log the report generation
        _log_report_activity(request, 'Program Summary Report', 'pdf', total_students, f'Period: {period_label}')
        
        return response

    except Exception as e:
        return HttpResponse(f"Error generating program summary: {str(e)}", status=500)


@coordinator_required
def generate_activity_log_report(request):
    """
    Generate Activity Log PDF showing coordinator actions.
    """
    import datetime
    from django.contrib.admin.models import LogEntry
    from django.contrib.contenttypes.models import ContentType

    try:
        program_obj, program_code, school_year = _get_coordinator_context(request)
        program_full = PROGRAM_NAMES.get(program_code, program_code)
        year_label = school_year.year_label if school_year else 'N/A'

        # Get recent log entries for this user (last 30 days)
        thirty_days_ago = datetime.datetime.now() - datetime.timedelta(days=30)
        
        # Try to get admin log entries for this user
        log_entries = LogEntry.objects.filter(
            user=request.user,
            action_time__gte=thirty_days_ago
        ).order_by('-action_time')[:100]

        # Generate PDF
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{program_code}_Activity_Log.pdf"'

        doc = SimpleDocTemplate(response, pagesize=landscape(letter))
        elements = []
        styles = getSampleStyleSheet()

        # Title
        elements.append(Paragraph(f"<b>{program_code} Coordinator Activity Log</b>", styles['Title']))
        elements.append(Spacer(1, 0.1 * inch))
        elements.append(Paragraph(f"Program: {program_full} | School Year: {year_label}", styles['Normal']))
        elements.append(Paragraph(f"Coordinator: {request.user.get_full_name() or request.user.username}", styles['Normal']))
        elements.append(Paragraph(f"Period: Last 30 Days", styles['Normal']))
        elements.append(Spacer(1, 0.3 * inch))

        # Activity Log Table
        elements.append(Paragraph("<b>Recent Activities</b>", styles['Heading2']))
        elements.append(Spacer(1, 0.1 * inch))

        if log_entries.exists():
            data = [['Date/Time', 'Action', 'Object', 'Details']]
            
            action_labels = {0: 'Added', 1: 'Changed', 2: 'Deleted'}
            
            for entry in log_entries:
                action = action_labels.get(entry.action_flag, 'Unknown')
                data.append([
                    entry.action_time.strftime('%Y-%m-%d %H:%M'),
                    action,
                    str(entry.content_type) if entry.content_type else 'N/A',
                    entry.object_repr[:40] + '...' if len(entry.object_repr) > 40 else entry.object_repr,
                ])

            table = Table(data, colWidths=[100, 80, 120, 250])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#991B1B')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#FEF2F2')),
            ]))
            elements.append(table)
        else:
            elements.append(Paragraph("No activity log entries found for the last 30 days.", styles['Normal']))
            elements.append(Spacer(1, 0.2 * inch))
            elements.append(Paragraph("<i>Note: Activity logging may not be enabled for all actions in the system.</i>", styles['Normal']))

        elements.append(Spacer(1, 0.5 * inch))
        elements.append(Paragraph(f"<b>Total Entries: {log_entries.count()}</b>", styles['Normal']))
        elements.append(Spacer(1, 0.2 * inch))
        elements.append(Paragraph(f"<i>Report generated on {datetime.datetime.now().strftime('%B %d, %Y at %I:%M %p')}</i>", styles['Normal']))

        doc.build(elements)
        
        # Log the report generation
        _log_report_activity(request, 'Activity Log Report', 'pdf', log_entries.count(), 'Last 30 Days')
        
        return response

    except Exception as e:
        return HttpResponse(f"Error generating activity log: {str(e)}", status=500)


@coordinator_required
def download_template(request):
    """
    Download blank templates and forms as Word documents.
    Designed to match the online pre-enrollment form layout.
    """
    import datetime
    from io import BytesIO
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from docx.shared import Inches, Cm, Twips
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    template_type = request.GET.get('template', 'enrollment')
    
    try:
        program_obj, program_code, school_year = _get_coordinator_context(request)
        program_full = PROGRAM_NAMES.get(program_code, program_code)
        year_label = school_year.year_label if school_year else 'N/A'

        # All templates are now Word documents
        doc_file = Document()
        
        # Set page margins for better layout
        sections = doc_file.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
        
        # ========== HELPER FUNCTIONS ==========
        
        def add_school_header(doc):
            """Add school header matching online form design."""
            # Header table with logo placeholder and school info
            header_table = doc.add_table(rows=1, cols=2)
            header_table.autofit = False
            
            # Logo cell (left)
            logo_cell = header_table.rows[0].cells[0]
            logo_cell.width = Cm(3)
            logo_para = logo_cell.paragraphs[0]
            logo_para.add_run('[SCHOOL LOGO]')
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # School info cell (right)
            info_cell = header_table.rows[0].cells[1]
            info_cell.width = Cm(14)
            
            # School name
            name_para = info_cell.paragraphs[0]
            name_run = name_para.add_run('Zamboanga National High School West')
            name_run.font.bold = True
            name_run.font.size = Pt(16)
            name_run.font.color.rgb = RGBColor(153, 27, 27)  # Maroon
            
            # Address
            addr_para = info_cell.add_paragraph()
            addr_run = addr_para.add_run('R.T. Lim Boulevard Zamboanga City, Philippines')
            addr_run.font.size = Pt(10)
            addr_run.font.color.rgb = RGBColor(75, 85, 99)
            
            # School ID
            id_para = info_cell.add_paragraph()
            id_run = id_para.add_run('School I.D: 303942')
            id_run.font.size = Pt(9)
            id_run.font.color.rgb = RGBColor(107, 114, 128)
            
            doc.add_paragraph('')  # Spacer
            
            # Add gray border line
            border_para = doc.add_paragraph()
            border_para.paragraph_format.space_after = Pt(0)
            border_run = border_para.add_run('─' * 80)
            border_run.font.color.rgb = RGBColor(209, 213, 219)
            border_run.font.size = Pt(6)
            
            doc.add_paragraph('')
        
        def add_form_title(doc, title_text, subtitle_text=''):
            """Add form title with border-top style like online form."""
            # Create title box using table with shading
            title_table = doc.add_table(rows=1, cols=1)
            title_table.autofit = False
            title_cell = title_table.rows[0].cells[0]
            title_cell.width = Cm(17)
            
            # Add top border (maroon line)
            tc = title_cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:top w:val="single" w:sz="24" w:color="991B1B"/>'
                '</w:tcBorders>'
            )
            tcPr.append(tcBorders)
            
            # Title
            title_para = title_cell.paragraphs[0]
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(title_text)
            title_run.font.bold = True
            title_run.font.size = Pt(20)
            title_run.font.color.rgb = RGBColor(153, 27, 27)
            
            # Subtitle
            if subtitle_text:
                sub_para = title_cell.add_paragraph()
                sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sub_run = sub_para.add_run(subtitle_text)
                sub_run.font.size = Pt(12)
                sub_run.font.color.rgb = RGBColor(55, 65, 81)
            
            doc.add_paragraph('')
        
        def add_section_box(doc, section_title, section_subtitle=''):
            """Add section heading with left maroon border like online form."""
            # Create section box using table
            section_table = doc.add_table(rows=1, cols=1)
            section_table.autofit = False
            section_cell = section_table.rows[0].cells[0]
            section_cell.width = Cm(17)
            
            # Add left border (maroon) and light gray background
            tc = section_cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:left w:val="single" w:sz="24" w:color="991B1B"/>'
                '<w:top w:val="nil"/>'
                '<w:right w:val="nil"/>'
                '<w:bottom w:val="nil"/>'
                '</w:tcBorders>'
            )
            tcPr.append(tcBorders)
            
            # Light gray background
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F9FAFB"/>')
            tcPr.append(shading)
            
            # Section title
            title_para = section_cell.paragraphs[0]
            title_run = title_para.add_run(section_title)
            title_run.font.bold = True
            title_run.font.size = Pt(14)
            title_run.font.color.rgb = RGBColor(153, 27, 27)
            
            # Section subtitle
            if section_subtitle:
                sub_para = section_cell.add_paragraph()
                sub_run = sub_para.add_run(section_subtitle)
                sub_run.font.size = Pt(9)
                sub_run.font.color.rgb = RGBColor(107, 114, 128)
            
            return section_cell
        
        def add_form_field_table(doc, fields, cols=2):
            """Add form fields in a grid layout like online form."""
            # Calculate rows needed
            import math
            rows_needed = math.ceil(len(fields) / cols)
            
            table = doc.add_table(rows=rows_needed, cols=cols * 2)  # label + input for each col
            table.autofit = False
            
            field_idx = 0
            for row_idx in range(rows_needed):
                for col_idx in range(cols):
                    if field_idx >= len(fields):
                        break
                    
                    label, placeholder = fields[field_idx]
                    
                    # Label cell
                    label_cell = table.rows[row_idx].cells[col_idx * 2]
                    label_cell.width = Cm(4)
                    label_para = label_cell.paragraphs[0]
                    label_run = label_para.add_run(label)
                    label_run.font.bold = True
                    label_run.font.size = Pt(9)
                    label_run.font.color.rgb = RGBColor(55, 65, 81)
                    
                    # Input cell
                    input_cell = table.rows[row_idx].cells[col_idx * 2 + 1]
                    input_cell.width = Cm(4.5)
                    input_para = input_cell.paragraphs[0]
                    
                    # Add border-bottom style for input
                    input_run = input_para.add_run(placeholder)
                    input_run.font.size = Pt(10)
                    input_run.font.color.rgb = RGBColor(156, 163, 175)
                    
                    field_idx += 1
            
            doc.add_paragraph('')
        
        def add_checkbox_group(doc, label, options):
            """Add checkbox group like online form."""
            para = doc.add_paragraph()
            label_run = para.add_run(label + ' ')
            label_run.font.bold = True
            label_run.font.size = Pt(10)
            label_run.font.color.rgb = RGBColor(55, 65, 81)
            
            options_run = para.add_run('    '.join([f'☐ {opt}' for opt in options]))
            options_run.font.size = Pt(10)
        
        def add_radio_group(doc, label, options):
            """Add radio button group like online form."""
            para = doc.add_paragraph()
            label_run = para.add_run(label + ' ')
            label_run.font.bold = True
            label_run.font.size = Pt(10)
            label_run.font.color.rgb = RGBColor(55, 65, 81)
            
            options_run = para.add_run('    '.join([f'○ {opt}' for opt in options]))
            options_run.font.size = Pt(10)
        
        def add_form_row(table, label, value_placeholder='____________________________'):
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = value_placeholder
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
        
        def add_section_heading(doc, text):
            """Legacy function for backwards compatibility."""
            heading = doc.add_paragraph()
            run = heading.add_run(text)
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(153, 27, 27)
            return heading

        if template_type == 'narrative':
            # Generate Narrative Report Template
            title = doc_file.add_heading('Narrative Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc_file.add_paragraph('')
            doc_file.add_paragraph(f"Program: {program_full}")
            doc_file.add_paragraph(f"School Year: {year_label}")
            doc_file.add_paragraph(f"Coordinator: ____________________________")
            doc_file.add_paragraph(f"Date: ____________________________")
            
            doc_file.add_paragraph('')
            doc_file.add_heading('I. Executive Summary', level=1)
            doc_file.add_paragraph('[Write a brief overview of the reporting period, key achievements, and highlights.]')
            doc_file.add_paragraph('')
            
            doc_file.add_heading('II. Enrollment Status', level=1)
            doc_file.add_paragraph('[Provide enrollment statistics, trends, and any notable changes.]')
            doc_file.add_paragraph('')
            
            doc_file.add_heading('III. Academic Performance', level=1)
            doc_file.add_paragraph('[Summarize academic performance, GWA trends, honor students, and at-risk interventions.]')
            doc_file.add_paragraph('')
            
            doc_file.add_heading('IV. Challenges and Issues', level=1)
            doc_file.add_paragraph('[Describe any challenges encountered and how they were addressed.]')
            doc_file.add_paragraph('')
            
            doc_file.add_heading('V. Recommendations', level=1)
            doc_file.add_paragraph('[Provide recommendations for improvement or future actions.]')
            doc_file.add_paragraph('')
            
            doc_file.add_heading('VI. Conclusion', level=1)
            doc_file.add_paragraph('[Summarize the report and outline next steps.]')
            
            doc_file.add_paragraph('')
            doc_file.add_paragraph('Prepared by: ____________________________')
            doc_file.add_paragraph('Signature: ____________________________')
            doc_file.add_paragraph('Date: ____________________________')
            
            filename = "Narrative_Report_Template.docx"

        elif template_type == 'enrollment':
            # ============================================================
            # COMPREHENSIVE ENROLLMENT FORM - Matches Online Pre-Enrollment
            # Includes: StudentData, FamilyData, SurveyData fields
            # Design matches enrollment_app templates (studentData.html, familyData.html)
            # ============================================================
            
            # School Header (like online form)
            add_school_header(doc_file)
            
            # Form Title with maroon top border
            add_form_title(doc_file, 'STUDENT DATA FORM', f'School Year: {year_label}')
            
            # ============ PHOTO PLACEHOLDER (right side in online form) ============
            photo_para = doc_file.add_paragraph()
            photo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            photo_run = photo_para.add_run('┌─────────────────┐\n│                                    │\n│     STUDENT 1x1      │\n│         PHOTO              │\n│                                    │\n└─────────────────┘')
            photo_run.font.size = Pt(8)
            photo_run.font.color.rgb = RGBColor(107, 114, 128)
            
            # LRN Field
            lrn_para = doc_file.add_paragraph()
            lrn_run = lrn_para.add_run('LRN Number ')
            lrn_run.font.bold = True
            lrn_run.font.size = Pt(10)
            lrn_run.font.color.rgb = RGBColor(55, 65, 81)
            lrn_para.add_run('*').font.color.rgb = RGBColor(239, 68, 68)
            doc_file.add_paragraph('_' * 50)
            
            doc_file.add_paragraph('')
            
            # Enrolling As checkboxes
            add_checkbox_group(doc_file, 'Enrolling As:', ['New Student', 'Transferee', 'Old Student'])
            
            doc_file.add_paragraph('')
            
            # PWD/SPED section with yellow highlight style
            pwd_box = add_section_box(doc_file, 'Person with Disability (PWD) or SPED?')
            add_radio_group(doc_file, '', ['Yes', 'No'])
            spec_para = doc_file.add_paragraph()
            spec_run = spec_para.add_run('If yes, please specify: ')
            spec_run.font.size = Pt(9)
            spec_para.add_run('_' * 50)
            
            doc_file.add_paragraph('')
            
            # Working Student section
            work_box = add_section_box(doc_file, 'Are you a working student?')
            add_radio_group(doc_file, '', ['Yes', 'No'])
            work_spec = doc_file.add_paragraph()
            work_spec_run = work_spec.add_run('If yes, please specify: ')
            work_spec_run.font.size = Pt(9)
            work_spec.add_run('_' * 50)
            
            doc_file.add_paragraph('')
            
            # ============ SECTION A: STUDENT'S INFORMATION ============
            add_section_box(doc_file, "A. Student's Information Data", "* Please fill in the complete and correct details")
            
            doc_file.add_paragraph('')
            
            # Name fields (3-column layout like online form)
            name_label = doc_file.add_paragraph()
            name_label_run = name_label.add_run('NAME *')
            name_label_run.font.bold = True
            name_label_run.font.size = Pt(10)
            
            name_table = doc_file.add_table(rows=2, cols=3)
            name_table.autofit = False
            # Row 1: Input lines
            name_table.rows[0].cells[0].text = '_' * 25
            name_table.rows[0].cells[1].text = '_' * 25
            name_table.rows[0].cells[2].text = '_' * 25
            # Row 2: Labels
            for cell, label in zip(name_table.rows[1].cells, ['Last Name', 'First Name', 'Middle Name']):
                para = cell.paragraphs[0]
                run = para.add_run(label) if not para.text else None
                if run:
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(107, 114, 128)
                else:
                    cell.text = label
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(8)
                            r.font.color.rgb = RGBColor(107, 114, 128)
            
            doc_file.add_paragraph('')
            
            # Grid fields (like online form layout)
            grid_table = doc_file.add_table(rows=6, cols=4)
            grid_data = [
                ('Gender', '☐ Male  ☐ Female  ☐ Other', 'Date of Birth', '____ / ____ / ________ Age: ___'),
                ('Place of Birth', '_' * 20, 'Religion', '_' * 20),
                ('Mother Tongue', '_' * 20, 'Ethnic Tribe', '_' * 20),
            ]
            
            for row_idx, (label1, val1, label2, val2) in enumerate(grid_data):
                cells = grid_table.rows[row_idx].cells
                # First pair
                cells[0].text = label1
                for p in cells[0].paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(9)
                cells[1].text = val1
                # Second pair
                cells[2].text = label2
                for p in cells[2].paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(9)
                cells[3].text = val2
            
            doc_file.add_paragraph('')
            
            # Address
            addr_label = doc_file.add_paragraph()
            addr_run = addr_label.add_run('Present Home Address')
            addr_run.font.bold = True
            addr_run.font.size = Pt(10)
            doc_file.add_paragraph('_' * 80)
            
            doc_file.add_paragraph('')
            
            # Previous School Information (nested box like online form)
            prev_box = doc_file.add_paragraph()
            prev_run = prev_box.add_run('Previous School Information')
            prev_run.font.bold = True
            prev_run.font.size = Pt(10)
            
            prev_table = doc_file.add_table(rows=3, cols=2)
            prev_data = [
                ('Name of Last School Attended:', '_' * 35),
                ('Previous Grade and Section:', '_' * 35),
                ('School Year Last Attended:', '_' * 35),
            ]
            for row_idx, (label, val) in enumerate(prev_data):
                prev_table.rows[row_idx].cells[0].text = label
                prev_table.rows[row_idx].cells[1].text = val
                for p in prev_table.rows[row_idx].cells[0].paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(9)
            
            doc_file.add_paragraph('')
            
            # ============ SECTION B: FATHER'S INFORMATION ============
            add_section_box(doc_file, "B. Father's Information", "(Optional)")
            
            doc_file.add_paragraph('')
            
            # Father's name (3-column)
            fname_label = doc_file.add_paragraph()
            fname_run = fname_label.add_run('NAME:')
            fname_run.font.bold = True
            
            fname_table = doc_file.add_table(rows=2, cols=3)
            fname_table.rows[0].cells[0].text = '_' * 25
            fname_table.rows[0].cells[1].text = '_' * 25
            fname_table.rows[0].cells[2].text = '_' * 25
            fname_table.rows[1].cells[0].text = 'Last Name'
            fname_table.rows[1].cells[1].text = 'First Name'
            fname_table.rows[1].cells[2].text = 'Middle Name'
            
            # Father's other fields
            fdob_table = doc_file.add_table(rows=1, cols=4)
            fdob_table.rows[0].cells[0].text = 'Date of Birth:'
            fdob_table.rows[0].cells[1].text = '____ / ____ / ________ Age: ___'
            fdob_table.rows[0].cells[2].text = 'Occupation:'
            fdob_table.rows[0].cells[3].text = '_' * 20
            
            faddr_para = doc_file.add_paragraph()
            faddr_para.add_run('Complete Home Address: ').font.bold = True
            faddr_para.add_run('_' * 60)
            
            fcontact_table = doc_file.add_table(rows=1, cols=4)
            fcontact_table.rows[0].cells[0].text = 'Contact Number:'
            fcontact_table.rows[0].cells[1].text = '_' * 20
            fcontact_table.rows[0].cells[2].text = 'Email Address:'
            fcontact_table.rows[0].cells[3].text = '_' * 25
            
            doc_file.add_paragraph('')
            
            # ============ SECTION C: MOTHER'S INFORMATION ============
            add_section_box(doc_file, "C. Mother's Information", "(Optional)")
            
            doc_file.add_paragraph('')
            
            # Mother's name (3-column)
            mname_label = doc_file.add_paragraph()
            mname_run = mname_label.add_run('NAME:')
            mname_run.font.bold = True
            
            mname_table = doc_file.add_table(rows=2, cols=3)
            mname_table.rows[0].cells[0].text = '_' * 25
            mname_table.rows[0].cells[1].text = '_' * 25
            mname_table.rows[0].cells[2].text = '_' * 25
            mname_table.rows[1].cells[0].text = 'Last Name'
            mname_table.rows[1].cells[1].text = 'First Name'
            mname_table.rows[1].cells[2].text = 'Middle Name'
            
            # Mother's other fields
            mdob_table = doc_file.add_table(rows=1, cols=4)
            mdob_table.rows[0].cells[0].text = 'Date of Birth:'
            mdob_table.rows[0].cells[1].text = '____ / ____ / ________ Age: ___'
            mdob_table.rows[0].cells[2].text = 'Occupation:'
            mdob_table.rows[0].cells[3].text = '_' * 20
            
            maddr_para = doc_file.add_paragraph()
            maddr_para.add_run('Complete Home Address: ').font.bold = True
            maddr_para.add_run('_' * 60)
            
            mcontact_table = doc_file.add_table(rows=1, cols=4)
            mcontact_table.rows[0].cells[0].text = 'Contact Number:'
            mcontact_table.rows[0].cells[1].text = '_' * 20
            mcontact_table.rows[0].cells[2].text = 'Email Address:'
            mcontact_table.rows[0].cells[3].text = '_' * 25
            
            doc_file.add_paragraph('')
            
            # ============ SECTION D: GUARDIAN INFORMATION ============
            add_section_box(doc_file, 'D. Guardian Information', '(If different from parents)')
            
            doc_file.add_paragraph('')
            
            add_radio_group(doc_file, 'Official Guardian:', ['Father', 'Mother', 'Other (specify below)'])
            
            # Guardian's name
            gname_label = doc_file.add_paragraph()
            gname_run = gname_label.add_run('NAME:')
            gname_run.font.bold = True
            
            gname_table = doc_file.add_table(rows=2, cols=3)
            gname_table.rows[0].cells[0].text = '_' * 25
            gname_table.rows[0].cells[1].text = '_' * 25
            gname_table.rows[0].cells[2].text = '_' * 25
            gname_table.rows[1].cells[0].text = 'Last Name'
            gname_table.rows[1].cells[1].text = 'First Name'
            gname_table.rows[1].cells[2].text = 'Middle Name'
            
            grelation_para = doc_file.add_paragraph()
            grelation_para.add_run('Relationship to Student: ').font.bold = True
            grelation_para.add_run('_' * 30)
            
            gcontact_table = doc_file.add_table(rows=1, cols=4)
            gcontact_table.rows[0].cells[0].text = 'Contact Number:'
            gcontact_table.rows[0].cells[1].text = '_' * 20
            gcontact_table.rows[0].cells[2].text = 'Email:'
            gcontact_table.rows[0].cells[3].text = '_' * 25
            
            gaddr_para = doc_file.add_paragraph()
            gaddr_para.add_run('Complete Address: ').font.bold = True
            gaddr_para.add_run('_' * 60)
            
            doc_file.add_paragraph('')
            
            # ============ SECTION E: SURVEY / NON-ACADEMIC DATA ============
            add_section_box(doc_file, 'E. Student Profile Survey', '(Non-Academic Information)')
            
            doc_file.add_paragraph('')
            
            # Learning Style
            ls_para = doc_file.add_paragraph()
            ls_run = ls_para.add_run('Preferred Learning Style: ')
            ls_run.font.bold = True
            ls_para.add_run('☐ Visual    ☐ Auditory    ☐ Kinesthetic    ☐ Reading/Writing')
            
            # Study Hours
            sh_para = doc_file.add_paragraph()
            sh_run = sh_para.add_run('Average Study Hours per Day: ')
            sh_run.font.bold = True
            sh_para.add_run('☐ Less than 1 hour    ☐ 1-2 hours    ☐ 2-3 hours    ☐ More than 3 hours')
            
            # Study Environment
            se_para = doc_file.add_paragraph()
            se_run = se_para.add_run('Study Environment: ')
            se_run.font.bold = True
            se_para.add_run('☐ Quiet room    ☐ Shared space    ☐ Library    ☐ Other')
            
            # Support
            sup_para = doc_file.add_paragraph()
            sup_run = sup_para.add_run('Schoolwork Support at Home: ')
            sup_run.font.bold = True
            sup_para.add_run('☐ Parents    ☐ Siblings    ☐ Tutor    ☐ Self-study    ☐ None')
            
            doc_file.add_paragraph('')
            
            # Enjoyed Subjects
            es_para = doc_file.add_paragraph()
            es_run = es_para.add_run('Most Enjoyed Subjects (check all that apply): ')
            es_run.font.bold = True
            doc_file.add_paragraph('☐ Mathematics    ☐ Science    ☐ English    ☐ Filipino    ☐ Social Studies')
            doc_file.add_paragraph('☐ TLE    ☐ MAPEH    ☐ Values Education    ☐ Other: ____________')
            
            doc_file.add_paragraph('')
            
            # Device Availability
            dev_para = doc_file.add_paragraph()
            dev_run = dev_para.add_run('Device Availability: ')
            dev_run.font.bold = True
            dev_para.add_run('☐ Smartphone    ☐ Tablet    ☐ Laptop/Computer    ☐ None')
            
            # Internet Access
            net_para = doc_file.add_paragraph()
            net_run = net_para.add_run('Internet Access: ')
            net_run.font.bold = True
            net_para.add_run('☐ Stable WiFi    ☐ Mobile Data    ☐ Limited/Intermittent    ☐ No Access')
            
            # Distance from School
            dist_para = doc_file.add_paragraph()
            dist_run = dist_para.add_run('Distance from School: ')
            dist_run.font.bold = True
            dist_para.add_run('☐ Less than 1km    ☐ 1-3 km    ☐ 3-5 km    ☐ More than 5 km')
            
            doc_file.add_paragraph('')
            
            # ============ SECTION F: DECLARATION ============
            add_section_box(doc_file, 'F. Declaration and Agreement')
            
            doc_file.add_paragraph('')
            
            decl_para = doc_file.add_paragraph()
            decl_para.add_run('I hereby certify that the information provided above is true and correct to the best of my knowledge. I understand that any false information may result in the cancellation of my enrollment.')
            decl_para.paragraph_format.space_after = Pt(12)
            
            agree_para = doc_file.add_paragraph()
            agree_para.add_run('☐ I agree to the terms and conditions of enrollment.')
            
            doc_file.add_paragraph('')
            doc_file.add_paragraph('')
            
            # Signatures (styled like online form buttons)
            sig_table = doc_file.add_table(rows=2, cols=2)
            sig_table.rows[0].cells[0].text = '________________________________'
            sig_table.rows[0].cells[1].text = '________________________________'
            sig_table.rows[1].cells[0].text = "Student's Signature / Date"
            sig_table.rows[1].cells[1].text = "Parent/Guardian's Signature / Date"
            
            filename = "Enrollment_Form.docx"

        elif template_type == 'transfer':
            # ============ SECTION TRANSFER FORM ============
            add_school_header(doc_file)
            add_form_title(doc_file, 'SECTION TRANSFER REQUEST FORM', f'School Year: {year_label}')
            
            add_section_box(doc_file, 'Student Details')
            
            doc_file.add_paragraph('')
            
            # Student info in grid layout
            student_table = doc_file.add_table(rows=3, cols=4)
            student_data = [
                ('Student Name:', '_' * 30, 'LRN:', '_' * 20),
                ('Grade Level:', '_' * 15, 'Current Section:', '_' * 15),
                ('Requested Section:', '_' * 15, 'Date of Request:', '____ / ____ / ____'),
            ]
            for row_idx, (l1, v1, l2, v2) in enumerate(student_data):
                cells = student_table.rows[row_idx].cells
                cells[0].text = l1
                cells[1].text = v1
                cells[2].text = l2
                cells[3].text = v2
                for p in cells[0].paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(9)
                for p in cells[2].paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(9)
            
            doc_file.add_paragraph('')
            
            add_section_box(doc_file, 'Reason for Transfer')
            
            doc_file.add_paragraph('')
            add_checkbox_group(doc_file, '', ['Schedule conflict', 'Health/Medical reason'])
            add_checkbox_group(doc_file, '', ['Academic adjustment', 'Personal/Family reason'])
            add_checkbox_group(doc_file, '', ['Other (please specify)'])
            
            doc_file.add_paragraph('')
            exp_label = doc_file.add_paragraph()
            exp_label.add_run('Detailed explanation:').font.bold = True
            doc_file.add_paragraph('_' * 80)
            doc_file.add_paragraph('_' * 80)
            doc_file.add_paragraph('_' * 80)
            
            doc_file.add_paragraph('')
            
            add_section_box(doc_file, 'Approval Signatures')
            
            doc_file.add_paragraph('')
            
            approval_table = doc_file.add_table(rows=4, cols=3)
            approval_table.style = 'Table Grid'
            approval_table.rows[0].cells[0].text = 'Role'
            approval_table.rows[0].cells[1].text = 'Signature'
            approval_table.rows[0].cells[2].text = 'Date'
            approval_table.rows[1].cells[0].text = 'Student'
            approval_table.rows[2].cells[0].text = 'Parent/Guardian'
            approval_table.rows[3].cells[0].text = 'Coordinator'
            
            # Make header row bold with maroon background
            for cell in approval_table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                # Add maroon background
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="991B1B"/>')
                tcPr.append(shading)
            
            filename = "Section_Transfer_Form.docx"

        elif template_type == 'consent':
            # ============ PARENT CONSENT FORM ============
            add_school_header(doc_file)
            add_form_title(doc_file, 'PARENT/GUARDIAN CONSENT FORM', f'School Year: {year_label}')
            
            add_section_box(doc_file, 'Student Information')
            
            doc_file.add_paragraph('')
            
            # Student info grid
            consent_student_table = doc_file.add_table(rows=2, cols=4)
            consent_student_table.rows[0].cells[0].text = 'Student Name:'
            consent_student_table.rows[0].cells[1].text = '_' * 30
            consent_student_table.rows[0].cells[2].text = 'LRN:'
            consent_student_table.rows[0].cells[3].text = '_' * 20
            consent_student_table.rows[1].cells[0].text = 'Grade & Section:'
            consent_student_table.rows[1].cells[1].text = '_' * 30
            
            for row in consent_student_table.rows:
                for idx, cell in enumerate(row.cells):
                    if idx % 2 == 0:  # Label cells
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.bold = True
                                r.font.size = Pt(9)
            
            doc_file.add_paragraph('')
            
            add_section_box(doc_file, 'Activity Details')
            
            doc_file.add_paragraph('')
            
            # Activity details in grid layout
            activity_table = doc_file.add_table(rows=3, cols=4)
            activity_data = [
                ('Activity/Event Name:', '_' * 30, 'Date:', '____ / ____ / ____'),
                ('Time:', '________ to ________', 'Venue/Location:', '_' * 20),
                ('Supervising Teacher:', '_' * 30, '', ''),
            ]
            for row_idx, (l1, v1, l2, v2) in enumerate(activity_data):
                cells = activity_table.rows[row_idx].cells
                cells[0].text = l1
                cells[1].text = v1
                cells[2].text = l2
                cells[3].text = v2
                for p in cells[0].paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(9)
                for p in cells[2].paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(9)
            
            doc_file.add_paragraph('')
            
            add_section_box(doc_file, 'Consent Statement')
            
            doc_file.add_paragraph('')
            consent_text = doc_file.add_paragraph()
            consent_text.add_run('I, the undersigned parent/legal guardian, hereby give my consent for my child/ward to participate in the above-mentioned activity.')
            
            doc_file.add_paragraph('')
            ack_label = doc_file.add_paragraph()
            ack_label.add_run('I understand and acknowledge the following:').font.bold = True
            
            add_checkbox_group(doc_file, '', ['I have read and understood the details of the activity.'])
            add_checkbox_group(doc_file, '', ['I authorize school personnel to make emergency medical decisions if I cannot be reached.'])
            add_checkbox_group(doc_file, '', ['I agree to pick up my child at the designated time and place.'])
            add_checkbox_group(doc_file, '', ['I understand the school will take reasonable precautions for student safety.'])
            
            doc_file.add_paragraph('')
            
            add_section_box(doc_file, 'Emergency Contact Information')
            
            doc_file.add_paragraph('')
            
            emergency_table = doc_file.add_table(rows=2, cols=4)
            emergency_table.rows[0].cells[0].text = 'Emergency Contact Number:'
            emergency_table.rows[0].cells[1].text = '_' * 25
            emergency_table.rows[0].cells[2].text = ''
            emergency_table.rows[0].cells[3].text = ''
            emergency_table.rows[1].cells[0].text = 'Alternative Contact Person:'
            emergency_table.rows[1].cells[1].text = '_' * 25
            emergency_table.rows[1].cells[2].text = 'Relationship:'
            emergency_table.rows[1].cells[3].text = '_' * 15
            
            for row in emergency_table.rows:
                for idx, cell in enumerate(row.cells):
                    if idx % 2 == 0:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.bold = True
                                r.font.size = Pt(9)
            
            doc_file.add_paragraph('')
            doc_file.add_paragraph('')
            
            sig_table = doc_file.add_table(rows=2, cols=2)
            sig_table.rows[0].cells[0].text = '________________________________'
            sig_table.rows[0].cells[1].text = '________________________________'
            sig_table.rows[1].cells[0].text = "Parent/Guardian's Signature"
            sig_table.rows[1].cells[1].text = 'Date'
            
            filename = "Parent_Consent_Form.docx"

        elif template_type == 'gradesheet':
            # ============ GRADE SHEET TEMPLATE ============
            add_school_header(doc_file)
            add_form_title(doc_file, 'CLASS GRADE SHEET', f'School Year: {year_label}')
            
            add_section_box(doc_file, 'Class Information')
            
            doc_file.add_paragraph('')
            
            # Class info in grid layout
            class_info_table = doc_file.add_table(rows=2, cols=4)
            class_info_data = [
                ('Section:', '_' * 20, 'Subject:', '_' * 20),
                ('Teacher:', '_' * 20, 'Quarter:', '☐ 1st  ☐ 2nd  ☐ 3rd  ☐ 4th'),
            ]
            for row_idx, (l1, v1, l2, v2) in enumerate(class_info_data):
                cells = class_info_table.rows[row_idx].cells
                cells[0].text = l1
                cells[1].text = v1
                cells[2].text = l2
                cells[3].text = v2
                for p in cells[0].paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(9)
                for p in cells[2].paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(9)
            
            doc_file.add_paragraph('')
            
            add_section_box(doc_file, 'Student Grades')
            
            doc_file.add_paragraph('')
            
            # Grade table
            grade_table = doc_file.add_table(rows=26, cols=7)
            grade_table.style = 'Table Grid'
            
            # Headers
            headers = ['#', 'Student Name', 'Q1', 'Q2', 'Q3', 'Q4', 'Final']
            for i, header in enumerate(headers):
                cell = grade_table.rows[0].cells[i]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                # Add maroon background to header
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="991B1B"/>')
                cell._tc.get_or_add_tcPr().append(shading_elm)
            
            # Empty rows for data
            for i in range(1, 26):
                grade_table.rows[i].cells[0].text = str(i)
            
            doc_file.add_paragraph('')
            doc_file.add_paragraph('')
            
            # Summary
            add_section_heading(doc_file, 'SUMMARY')
            summary_table = doc_file.add_table(rows=0, cols=2)
            summary_table.style = 'Table Grid'
            add_form_row(summary_table, 'Total Students:', '________')
            add_form_row(summary_table, 'Passing:', '________')
            add_form_row(summary_table, 'Failing:', '________')
            add_form_row(summary_table, 'Class Average:', '________')
            
            doc_file.add_paragraph('')
            doc_file.add_paragraph('')
            doc_file.add_paragraph('Prepared by: ________________________________')
            doc_file.add_paragraph('Date: ________________________________')
            
            filename = "Grade_Sheet_Template.docx"
        
        else:
            # Default fallback
            doc_file.add_heading('Template', 0)
            doc_file.add_paragraph('Template not found.')
            filename = "Template.docx"

        # Return Word document response
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc_file.save(response)
        
        # Log the template download
        _log_report_activity(request, f'Template: {template_type.title()}', 'word', 1, f'Template: {filename}')
        
        return response

    except Exception as e:
        return HttpResponse(f"Error generating template: {str(e)}", status=500)


# ──────────────────────────────────────────────
# Activity Log API
# ──────────────────────────────────────────────

from django.http import JsonResponse

@coordinator_required
def get_activity_logs(request):
    """
    API endpoint to get activity logs for the current coordinator.
    Returns logs filtered by the coordinator's program.
    """
    try:
        program_obj, program_code, school_year = _get_coordinator_context(request)
        
        # Get query parameters
        limit = int(request.GET.get('limit', 50))
        offset = int(request.GET.get('offset', 0))
        category = request.GET.get('category', 'all')
        
        # Limit max results
        limit = min(limit, 100)
        
        # Build queryset
        logs = CoordinatorActivityLog.objects.all()
        
        # Filter by coordinator's program
        if program_obj:
            logs = logs.filter(
                Q(program=program_obj) | Q(user=request.user)
            )
        else:
            logs = logs.filter(user=request.user)
        
        # Filter by category
        if category and category != 'all':
            logs = logs.filter(category=category)
        
        # Order by most recent first
        logs = logs.order_by('-created_at')
        
        # Get total count before pagination
        total_count = logs.count()
        
        # Apply pagination
        logs = logs[offset:offset + limit]
        
        # Format response data
        logs_data = []
        for log in logs:
            logs_data.append({
                'id': log.id,
                'action': log.action,
                'action_display': log.get_action_display(),
                'category': log.category,
                'category_display': log.get_category_display(),
                'description': log.description,
                'student_lrn': log.student_lrn,
                'student_name': log.student_name,
                'section_name': log.section_name,
                'user': log.user.get_full_name() if log.user else 'System',
                'date': log.get_formatted_date(),
                'time': log.get_formatted_time(),
                'icon_class': log.get_icon_class(),
                'color_class': log.get_color_class(),
                'created_at': log.created_at.isoformat(),
            })
        
        return JsonResponse({
            'success': True,
            'logs': logs_data,
            'total_count': total_count,
            'has_more': offset + limit < total_count,
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


@coordinator_required
def generate_custom_report(request):
    """Generate a custom report with user-selected data columns and filters."""
    from datetime import timedelta
    from django.utils import timezone
    
    try:
        program_obj, program_code, school_year = _get_coordinator_context(request)
        
        # Get parameters from request
        output_format = request.GET.get('format', 'pdf').lower()
        date_range = request.GET.get('date_range', 'Current School Year')
        
        # Get data inclusion options (checkboxes)
        include_demographics = request.GET.get('include_demographics', 'true') == 'true'
        include_grades = request.GET.get('include_grades', 'true') == 'true'
        include_sections = request.GET.get('include_sections', 'true') == 'true'
        include_contact = request.GET.get('include_contact', 'false') == 'true'
        
        year_label = school_year.year_label if school_year else 'N/A'
        program_full = PROGRAM_NAMES.get(program_code, program_code)
        
        # Get base selections
        active_grade = request.session.get('active_grade_code')
        selections = _get_base_selections(program_code, school_year, active_grade)
        
        # Apply date range filter if applicable
        now = timezone.now()
        if date_range == 'Last 30 Days':
            date_from = now - timedelta(days=30)
            selections = selections.filter(created_at__gte=date_from)
        elif date_range == 'Last Quarter':
            date_from = now - timedelta(days=90)
            selections = selections.filter(created_at__gte=date_from)
        # Current School Year and Custom Range use the full school year data
        
        # Build sections map for name lookup
        sections_map = {}
        if program_obj and school_year:
            for s in Section.objects.filter(program=program_obj, school_year=school_year):
                sections_map[str(s.id)] = s.name
        
        # Collect data based on selected columns
        student_rows = []
        headers = ['#']
        
        # Build headers based on options
        if include_demographics:
            headers.extend(['LRN', 'Last Name', 'First Name', 'Middle Name', 'Gender', 'Birth Date'])
        if include_grades:
            headers.extend(['Math', 'Science', 'English', 'Filipino', 'AP', 'ESP', 'TLE', 'MAPEH', 'GWA'])
        if include_sections:
            headers.extend(['Section', 'Status'])
        if include_contact:
            headers.extend(['Email', 'Phone', 'Address'])
        
        # Collect data rows
        for idx, ps in enumerate(selections, 1):
            student_info = getattr(ps.student, 'student_data', None)
            academic = getattr(ps.student, 'academic_data', None)
            
            if student_info:
                row = {'num': idx}
                
                if include_demographics:
                    row.update({
                        'lrn': ps.student.lrn,
                        'last_name': student_info.last_name or '',
                        'first_name': student_info.first_name or '',
                        'middle_name': student_info.middle_name or '',
                        'gender': (student_info.gender or '').capitalize(),
                        'birth_date': student_info.date_of_birth.strftime('%Y-%m-%d') if student_info.date_of_birth else '',
                    })
                
                if include_grades and academic:
                    row.update({
                        'math': str(academic.mathematics or ''),
                        'science': str(academic.science or ''),
                        'english': str(academic.english or ''),
                        'filipino': str(academic.filipino or ''),
                        'ap': str(academic.araling_panlipunan or ''),
                        'esp': str(academic.edukasyon_sa_pagpapakatao or ''),
                        'tle': str(academic.edukasyon_pangkabuhayan or ''),
                        'mapeh': str(academic.mapeh or ''),
                        'gwa': str(_calculate_gwa(academic)),
                    })
                elif include_grades:
                    row.update({k: '' for k in ['math', 'science', 'english', 'filipino', 'ap', 'esp', 'tle', 'mapeh', 'gwa']})
                
                if include_sections:
                    row.update({
                        'section': _get_section_name(ps, sections_map),
                        'status': _get_enrollment_status(ps),
                    })
                
                if include_contact:
                    # Email is on Student model, not StudentData
                    email = ps.student.email if ps.student.email else ''
                    # Try to get contact from family data (guardian)
                    phone = ''
                    family = getattr(ps.student, 'family_data', None)
                    if family:
                        phone = family.official_guardian_contact or ''
                    # Address is on StudentData
                    address = student_info.address if hasattr(student_info, 'address') and student_info.address else ''
                    row.update({
                        'email': email or '',
                        'phone': phone or '',
                        'address': address,
                    })
                
                student_rows.append(row)
        
        # Map for row data keys to match headers
        key_order = ['num']
        if include_demographics:
            key_order.extend(['lrn', 'last_name', 'first_name', 'middle_name', 'gender', 'birth_date'])
        if include_grades:
            key_order.extend(['math', 'science', 'english', 'filipino', 'ap', 'esp', 'tle', 'mapeh', 'gwa'])
        if include_sections:
            key_order.extend(['section', 'status'])
        if include_contact:
            key_order.extend(['email', 'phone', 'address'])
        
        # Generate report in requested format
        if output_format == 'excel':
            return _generate_custom_excel(student_rows, headers, key_order, program_code, program_full, year_label, date_range)
        elif output_format == 'word':
            return _generate_custom_word(student_rows, headers, key_order, program_code, program_full, year_label, date_range)
        else:
            return _generate_custom_pdf(student_rows, headers, key_order, program_code, program_full, year_label, date_range)
        
        # Log the activity
        _log_report_activity(request, 'Custom Report', output_format, len(student_rows), f'Date Range: {date_range}')
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return HttpResponse(f'Error generating custom report: {str(e)}', status=500)


def _generate_custom_excel(student_rows, headers, key_order, program_code, program_full, year_label, date_range):
    """Generate custom report as Excel file."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Custom Report"
    
    # Title
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
    ws.cell(row=1, column=1).value = f"{program_code} Custom Report"
    ws.cell(row=1, column=1).font = Font(bold=True, size=16, color="991B1B")
    ws.cell(row=1, column=1).alignment = Alignment(horizontal='center')
    
    # Subtitle
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(headers))
    ws.cell(row=2, column=1).value = f"Program: {program_full} | School Year: {year_label} | Date Range: {date_range}"
    ws.cell(row=2, column=1).alignment = Alignment(horizontal='center')
    
    # Headers
    header_fill = PatternFill(start_color="991B1B", end_color="991B1B", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col_num)
        cell.value = header
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
    
    # Data rows
    for row_idx, row_data in enumerate(student_rows, 5):
        for col_idx, key in enumerate(key_order, 1):
            ws.cell(row=row_idx, column=col_idx).value = row_data.get(key, '')
    
    # Auto-adjust column widths (skip merged cells)
    from openpyxl.utils import get_column_letter
    for col_idx in range(1, len(headers) + 1):
        column_letter = get_column_letter(col_idx)
        max_length = 0
        for row_idx in range(4, len(student_rows) + 5):  # Start from header row (4)
            cell = ws.cell(row=row_idx, column=col_idx)
            try:
                if cell.value and len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        ws.column_dimensions[column_letter].width = min(max_length + 2, 30)
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename={program_code}_custom_report.xlsx'
    wb.save(response)
    return response


def _generate_custom_word(student_rows, headers, key_order, program_code, program_full, year_label, date_range):
    """Generate custom report as Word document."""
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from docx.shared import Inches
    
    doc = Document()
    
    # Title
    title = doc.add_heading(f'{program_code} Custom Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    info_para = doc.add_paragraph()
    info_run = info_para.add_run(f"Program: {program_full} | School Year: {year_label} | Date Range: {date_range}")
    info_run.bold = True
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    # Create table with narrower columns if many columns
    num_cols = len(headers)
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(8 if num_cols > 10 else 10)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="991B1B"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    
    # Data rows
    for row_data in student_rows:
        row_cells = table.add_row().cells
        for col_idx, key in enumerate(key_order):
            row_cells[col_idx].text = str(row_data.get(key, ''))
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8 if num_cols > 10 else 9)
    
    # Summary
    doc.add_paragraph('')
    summary = doc.add_paragraph()
    summary_run = summary.add_run(f"Total Records: {len(student_rows)}")
    summary_run.bold = True
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename={program_code}_custom_report.docx'
    doc.save(response)
    return response


def _generate_custom_pdf(student_rows, headers, key_order, program_code, program_full, year_label, date_range):
    """Generate custom report as PDF."""
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename={program_code}_custom_report.pdf'
    
    # Use landscape if many columns
    num_cols = len(headers)
    page_size = landscape(letter) if num_cols > 8 else letter
    
    doc = SimpleDocTemplate(response, pagesize=page_size, topMargin=0.5*inch, bottomMargin=0.5*inch)
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = styles['Heading1']
    title_style.alignment = 1  # Center
    elements.append(Paragraph(f'{program_code} Custom Report', title_style))
    
    # Subtitle
    subtitle = f"Program: {program_full} | School Year: {year_label} | Date Range: {date_range}"
    elements.append(Paragraph(subtitle, styles['Normal']))
    elements.append(Spacer(1, 0.25*inch))
    
    # Build table data
    table_data = [headers]
    for row_data in student_rows:
        row = [str(row_data.get(key, '')) for key in key_order]
        table_data.append(row)
    
    # Calculate column widths
    available_width = page_size[0] - 1*inch  # Page width minus margins
    col_width = available_width / num_cols
    
    # Create table
    table = Table(table_data, colWidths=[col_width] * num_cols)
    
    # Style table
    style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.6, 0.1, 0.1)),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 7 if num_cols > 10 else 9),
        ('FONTSIZE', (0, 1), (-1, -1), 6 if num_cols > 10 else 8),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.Color(0.95, 0.95, 0.95)]),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
    ])
    table.setStyle(style)
    
    elements.append(table)
    elements.append(Spacer(1, 0.25*inch))
    elements.append(Paragraph(f"<b>Total Records:</b> {len(student_rows)}", styles['Normal']))
    
    doc.build(elements)
    return response
